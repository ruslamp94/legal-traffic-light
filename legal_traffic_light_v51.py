"""
Legal Traffic Light v5.1 - Enterprise Edition
Система анализа договоров на соответствие Регламенту АО «НПК»

Функционал:
- Загрузка документов: TXT, DOCX, PDF
- Генерация отчетов: PDF, DOCX, JSON
- Алгоритм сравнения: Jaccard + TF-IDF + N-grams + Levenshtein
- AI-анализ рисков (OpenAI/Anthropic)
- История анализов, статус согласования ЮД

Запуск: streamlit run app.py
Зависимости: pip install streamlit python-docx fpdf2 PyPDF2 requests
"""

import streamlit as st
import re
import json
import html
import math
import hashlib
import base64
import io
import os
from datetime import datetime, date
from typing import Dict, List, Tuple, Optional, Any, Set
from dataclasses import dataclass, field, asdict
from enum import Enum
from collections import Counter
import difflib

# ============================================================================
# КОНФИГУРАЦИЯ
# ============================================================================

st.set_page_config(
    page_title="Legal Traffic Light v5.1 | АО НПК",
    page_icon="🚦",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# КОНСТАНТЫ ИЗ РЕГЛАМЕНТА
# ============================================================================

class RiskZone(Enum):
    GREEN = "green"
    YELLOW = "yellow"
    RED = "red"

class DocumentForm(Enum):
    TYPICAL = "typical"
    COUNTERPARTY = "counterparty"
    FREE = "free"
    MODIFIED_TF = "modified_tf"
    SELF_DEVELOPED = "self"

class LegalStatus(Enum):
    NOT_SUBMITTED = "not_submitted"
    NO_INFO = "no_info"
    SUBMITTED = "submitted"
    APPROVED = "approved"
    REJECTED = "rejected"
    IN_PROGRESS = "in_progress"

LEGAL_STATUS_LABELS = {
    LegalStatus.NOT_SUBMITTED: "🔴 Не поступал на согласование ЮД",
    LegalStatus.NO_INFO: "🟡 Нет сведений о поступлении в данной редакции",
    LegalStatus.SUBMITTED: "🟡 Поступал на согласование",
    LegalStatus.APPROVED: "🟢 Согласован ЮД",
    LegalStatus.REJECTED: "🔴 Отклонен ЮД",
    LegalStatus.IN_PROGRESS: "🟡 На рассмотрении ЮД",
}

class Thresholds:
    GREEN_TF_MAX = 100_000
    GREEN_NON_TF_MAX = 50_000
    YELLOW_MAX = 5_000_000
    RED_MIN = 5_000_001
    TENDER_RED = 3_000_000
    SINGLE_SUPPLIER_YELLOW = 100_000
    CONTRACT_CONTROL_YEARS = 3

class Deadlines:
    STANDARD = 5
    EXTENDED = 10
    URGENT = 1

# Подразделения АО НПК
DEPARTMENTS = [
    "Департамент организации перевозок (универсальный подвижной состав)",
    "Департамент по инвестициям",
    "Департамент по расчетам с ОАО \"РЖД\"",
    "Департамент по связям с общественностью",
    "Департамент по управлению персоналом",
    "Департамент подвижного состава",
    "Департамент развития персонала",
    "Департамент стратегических проектов",
    "Департамент финансового контроллинга",
    "Казначейство",
    "Юридический департамент",
    "Отдел по управлению рисками",
    "Служба охраны труда",
]

POSITIONS = [
    "Специалист",
    "Ведущий специалист",
    "Главный специалист",
    "Начальник отдела",
    "Заместитель руководителя управления",
    "Начальник управления",
    "Заместитель руководителя департамента",
    "Руководитель департамента",
    "Заместитель генерального директора",
    "Генеральный директор"
]

RED_ZONE_ALWAYS = [
    "Аренда вагонов", "Лизинг вагонов", "Покупка/продажа вагонов",
    "Аренда локомотивов", "Лизинг локомотивов", "Покупка/продажа локомотивов",
    "Аренда контейнеров", "Лизинг контейнеров", "Покупка/продажа контейнеров",
    "Международные перевозки (ВЭД)", "Расчеты в валюте",
    "Договор на разработку ПО", "Договор на внедрение ПО",
    "Лицензионное соглашение на ПО", "Смарт-контракт",
    "Аренда недвижимости", "Покупка недвижимости",
    "Кредитный договор", "Договор займа", "Договор залога", "Договор поручительства",
    "Договор с ОАО РЖД", "Сервисный (глобальный) договор",
    "Договор, требующий одобрения Совета директоров",
    "Трудовой договор с ТОП-менеджментом",
    "Локальный нормативный акт (ЛНА)", "Положение/Правила/Инструкция",
    "Приказ о коммерческой тайне", "Приказ о дисциплинарном взыскании",
    "Приказ о материальной ответственности",
]

YELLOW_ZONE_TYPES = [
    "Договор на регулярные (рамочные) перевозки",
    "Договор на годовые перевозки",
    "Договор ТЭУ (транспортно-экспедиционные услуги)",
    "Перевозка опасного груза", "Перевозка негабаритного груза",
    "Перевозка тяжеловесного груза", "Перевозка дорогостоящего груза",
    "Закупка у единственного поставщика",
]

RED_DOCUMENTS_ALWAYS = [
    "Претензия (входящая)", "Претензия (исходящая)",
    "Исковое заявление", "Судебный приказ",
    "Запрос ФНС", "Запрос ФАС", "Запрос Прокуратуры",
    "Запрос Ространснадзора", "Запрос ГИТ (трудовая инспекция)",
    "Запрос иного госоргана", "Предписание госоргана", "Требование госоргана",
    "ДТП с участием ТС компании", "Утеря груза", "Порча груза",
    "Простой, требующий юридической фиксации",
]

# ============================================================================
# ТИПОВЫЕ ФОРМЫ (ВСТРОЕННЫЕ)
# ============================================================================

BUILTIN_TYPICAL_FORMS = {
    "service": {
        "name": "Типовая форма договора оказания услуг",
        "code": "ТФ-УСЛ-001",
        "version": "3.0",
        "date": "01.01.2025",
        "sections": {
            "1. ПРЕДМЕТ ДОГОВОРА": {
                "required": True,
                "template": """1.1. Исполнитель обязуется оказать Заказчику услуги, указанные в Техническом задании (Приложение №1), а Заказчик обязуется принять и оплатить оказанные услуги.""",
                "keywords": ["предмет", "услуги", "обязуется", "оказать", "принять", "оплатить", "техническое задание"],
                "risk_patterns": []
            },
            "2. СТОИМОСТЬ И ПОРЯДОК РАСЧЕТОВ": {
                "required": True,
                "template": """2.1. Стоимость услуг составляет _______ рублей. 2.2. Оплата производится в течение 10 рабочих дней с даты подписания Акта.""",
                "keywords": ["стоимость", "оплата", "рабочих дней", "акт", "ндс", "расчет"],
                "risk_patterns": [
                    {"pattern": r"предоплат.*(?:[3-9]\d|100)\s*%", "risk": "red", "issue": "Предоплата более 30%"},
                    {"pattern": r"оплата.*(?:1|2|3)\s*(?:рабоч|календарн)", "risk": "yellow", "issue": "Слишком короткий срок оплаты"},
                ]
            },
            "3. СРОКИ ОКАЗАНИЯ УСЛУГ": {
                "required": True,
                "template": """3.1. Срок оказания услуг: с «___» ________ 202__ г. по «___» ________ 202__ г.""",
                "keywords": ["срок", "оказания", "услуг", "период"],
                "risk_patterns": []
            },
            "4. ПОРЯДОК СДАЧИ-ПРИЕМКИ УСЛУГ": {
                "required": True,
                "template": """4.1. По завершении оказания услуг Исполнитель направляет Заказчику Акт в 2-х экземплярах. 4.2. Заказчик в течение 5 рабочих дней обязан подписать его или направить мотивированный отказ.""",
                "keywords": ["приемка", "акт", "сдача", "рабочих дней", "мотивированный отказ"],
                "risk_patterns": [
                    {"pattern": r"(?:1|2)\s*(?:рабоч|календарн).*(?:подпис|приня)", "risk": "yellow", "issue": "Слишком короткий срок приемки"},
                ]
            },
            "5. ПРАВА И ОБЯЗАННОСТИ СТОРОН": {
                "required": True,
                "template": """5.1. Исполнитель обязан оказать услуги надлежащего качества. 5.2. Заказчик обязан своевременно оплатить услуги.""",
                "keywords": ["права", "обязанности", "исполнитель", "заказчик", "обязан"],
                "risk_patterns": []
            },
            "6. ОТВЕТСТВЕННОСТЬ СТОРОН": {
                "required": True,
                "template": """6.1. За нарушение сроков оплаты неустойка 0,1% за каждый день просрочки, но не более 10%. 6.2. За нарушение сроков оказания услуг неустойка 0,1% за каждый день просрочки, но не более 10%.""",
                "keywords": ["ответственность", "неустойка", "просрочка", "штраф"],
                "risk_patterns": [
                    {"pattern": r"ответственност.*ограничен.*(?:последн|месяц|платеж)", "risk": "red", "issue": "Ограничение ответственности суммой последнего платежа"},
                    {"pattern": r"не\s+(?:несет|отвечает).*(?:косвенн|упущенн)", "risk": "red", "issue": "Исключение косвенных убытков"},
                    {"pattern": r"неустойк.*(?:0[,.]?[5-9]|[1-9][,.]?\d)\s*%", "risk": "red", "issue": "Неустойка выше 0.3% в день"},
                    {"pattern": r"(?:полн\w+|неограничен\w+).*ответственност", "risk": "red", "issue": "Неограниченная ответственность"},
                ]
            },
            "7. КОНФИДЕНЦИАЛЬНОСТЬ": {
                "required": True,
                "template": """7.1. Стороны обязуются не разглашать конфиденциальную информацию в течение 3 лет после прекращения Договора.""",
                "keywords": ["конфиденциальность", "разглашать", "информация", "секрет"],
                "risk_patterns": [
                    {"pattern": r"(?:штраф|неустойк).*конфиденциальност.*(?:[5-9]|1\d)\s*(?:000\s*000|млн)", "risk": "red", "issue": "Чрезмерный штраф за конфиденциальность (>5 млн)"},
                ]
            },
            "8. ИНТЕЛЛЕКТУАЛЬНАЯ СОБСТВЕННОСТЬ": {
                "required": False,
                "template": """8.1. Исключительные права на результаты переходят к Заказчику с момента подписания Акта.""",
                "keywords": ["интеллектуальная", "собственность", "права", "результат", "исключительные"],
                "risk_patterns": [
                    {"pattern": r"(?:исключительн\w+\s+)?прав\w+.*(?:принадлежат?|остают?ся|переходят?).*исполнител", "risk": "red", "issue": "Права на результаты остаются у Исполнителя"},
                    {"pattern": r"неисключительн\w+\s+лицензи", "risk": "yellow", "issue": "Только неисключительная лицензия для Заказчика"},
                ]
            },
            "9. СРОК ДЕЙСТВИЯ И РАСТОРЖЕНИЕ": {
                "required": True,
                "template": """9.1. Договор действует до полного исполнения обязательств. 9.2. Любая Сторона вправе расторгнуть Договор с уведомлением за 30 дней.""",
                "keywords": ["срок", "действия", "расторжение", "уведомление", "односторонн"],
                "risk_patterns": [
                    {"pattern": r"исполнитель.*(?:вправе|имеет\s+право).*односторонн.*расторг.*(?:[1-9]|1[0-4])\s*(?:дн|календарн)", "risk": "red", "issue": "Односторонний отказ Исполнителя с коротким сроком (<15 дней)"},
                    {"pattern": r"заказчик.*только.*(?:существенн|нарушен)", "risk": "red", "issue": "Ограничение права Заказчика на расторжение"},
                    {"pattern": r"автоматическ\w+\s+(?:пролонгац|продлен)", "risk": "yellow", "issue": "Автоматическая пролонгация"},
                ]
            },
            "10. РАЗРЕШЕНИЕ СПОРОВ": {
                "required": True,
                "template": """10.1. Споры разрешаются путем переговоров. 10.2. При недостижении согласия — в Арбитражном суде г. Москвы.""",
                "keywords": ["споры", "арбитражный", "суд", "разногласия", "переговоры"],
                "risk_patterns": [
                    {"pattern": r"арбитражн\w+\s+суд\w*.*(?:санкт-петербург|спб|питер)", "risk": "yellow", "issue": "Подсудность в Санкт-Петербурге"},
                    {"pattern": r"третейск\w+\s+суд", "risk": "yellow", "issue": "Третейская оговорка"},
                ]
            },
            "11. ФОРС-МАЖОР": {
                "required": True,
                "template": """11.1. Стороны освобождаются от ответственности при обстоятельствах непреодолимой силы.""",
                "keywords": ["форс-мажор", "непреодолимой силы", "обстоятельства"],
                "risk_patterns": []
            },
            "12. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ": {
                "required": True,
                "template": """12.1. Договор составлен в двух экземплярах. 12.2. Приложения являются неотъемлемой частью Договора.""",
                "keywords": ["заключительные", "экземпляр", "приложения", "изменения"],
                "risk_patterns": []
            },
            "13. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН": {
                "required": True,
                "template": """ЗАКАЗЧИК: АО «НПК»... ИСПОЛНИТЕЛЬ: ...""",
                "keywords": ["реквизиты", "подписи", "заказчик", "исполнитель", "инн", "адрес"],
                "risk_patterns": []
            }
        },
        "global_risk_patterns": [
            {"pattern": r"односторонн\w+.*изменен\w+.*(?:цен|стоимост|тариф)", "risk": "red", "issue": "Одностороннее изменение цены контрагентом"},
            {"pattern": r"субподряд.*без.*(?:согласи|уведомлени)", "risk": "yellow", "issue": "Субподряд без согласия Заказчика"},
            {"pattern": r"(?:usd|eur|евро|доллар|валют|у\.е\.)", "risk": "yellow", "issue": "Валютная оговорка"},
        ]
    },
    "supply": {
        "name": "Типовая форма договора поставки",
        "code": "ТФ-ПОС-001",
        "version": "2.0",
        "date": "01.01.2025",
        "sections": {
            "1. ПРЕДМЕТ ДОГОВОРА": {"required": True, "template": "Поставщик обязуется передать товар.", "keywords": ["предмет", "поставщик", "покупатель", "товар"], "risk_patterns": []},
            "2. КАЧЕСТВО И КОМПЛЕКТНОСТЬ": {"required": True, "template": "Качество по ГОСТ.", "keywords": ["качество", "комплектность", "гост"], "risk_patterns": []},
            "3. ЦЕНА И РАСЧЕТЫ": {"required": True, "template": "Цена в Спецификации.", "keywords": ["цена", "оплата", "расчет"], "risk_patterns": [{"pattern": r"предоплат.*(?:[3-9]\d|100)\s*%", "risk": "red", "issue": "Предоплата более 30%"}]},
            "4. СРОКИ ПОСТАВКИ": {"required": True, "template": "Поставка по графику.", "keywords": ["срок", "поставка"], "risk_patterns": []},
            "5. ПОРЯДОК ПРИЕМКИ": {"required": True, "template": "Приемка по П-6, П-7.", "keywords": ["приемка", "п-6", "п-7"], "risk_patterns": []},
            "6. ОТВЕТСТВЕННОСТЬ": {"required": True, "template": "Неустойка 0,1%.", "keywords": ["ответственность", "неустойка"], "risk_patterns": []},
            "7. ГАРАНТИИ": {"required": True, "template": "Гарантийный срок.", "keywords": ["гарантия"], "risk_patterns": []},
            "8. ФОРС-МАЖОР": {"required": True, "template": "Непреодолимая сила.", "keywords": ["форс-мажор"], "risk_patterns": []},
            "9. СПОРЫ": {"required": True, "template": "Арбитраж Москвы.", "keywords": ["споры", "арбитражный"], "risk_patterns": []},
            "10. РЕКВИЗИТЫ": {"required": True, "template": "Реквизиты сторон.", "keywords": ["реквизиты", "подписи"], "risk_patterns": []}
        },
        "global_risk_patterns": [{"pattern": r"переход.*риск.*до.*передач", "risk": "red", "issue": "Переход рисков до передачи товара"}]
    },
    "wagon_rent": {
        "name": "Типовая форма договора аренды вагонов",
        "code": "ТФ-АРВ-001",
        "version": "2.0",
        "date": "01.01.2025",
        "sections": {
            "1. ПРЕДМЕТ ДОГОВОРА": {"required": True, "template": "Аренда вагонов.", "keywords": ["предмет", "арендодатель", "арендатор", "вагон"], "risk_patterns": []},
            "2. ПЕРЕЧЕНЬ ВАГОНОВ": {"required": True, "template": "Список вагонов.", "keywords": ["перечень", "номер", "вагон"], "risk_patterns": []},
            "3. АРЕНДНАЯ ПЛАТА": {"required": True, "template": "Ставка за сутки.", "keywords": ["аренда", "ставка", "плата"], "risk_patterns": []},
            "4. ПЕРЕДАЧА И ВОЗВРАТ": {"required": True, "template": "Акт передачи.", "keywords": ["передача", "возврат", "акт"], "risk_patterns": []},
            "5. ОБЯЗАННОСТИ АРЕНДАТОРА": {"required": True, "template": "Использование по назначению.", "keywords": ["обязанности", "арендатор"], "risk_patterns": []},
            "6. ОТВЕТСТВЕННОСТЬ": {"required": True, "template": "Ответственность за утрату.", "keywords": ["ответственность", "сохранность", "утрата"], "risk_patterns": [
                {"pattern": r"(?:полн\w+|неограничен\w+).*ответственност.*утрат", "risk": "red", "issue": "Неограниченная ответственность за утрату"},
                {"pattern": r"штраф.*простой.*(?:[2-9]\d{3}|[1-9]\d{4})", "risk": "red", "issue": "Штраф за простой более 2000 руб/сутки"}
            ]},
            "7. СТРАХОВАНИЕ": {"required": True, "template": "Страховка.", "keywords": ["страхование"], "risk_patterns": []},
            "8. СРОК АРЕНДЫ": {"required": True, "template": "Срок аренды.", "keywords": ["срок", "аренда"], "risk_patterns": []},
            "9. СПОРЫ": {"required": True, "template": "Арбитраж.", "keywords": ["споры", "арбитражный"], "risk_patterns": []},
            "10. РЕКВИЗИТЫ": {"required": True, "template": "Реквизиты.", "keywords": ["реквизиты", "подписи"], "risk_patterns": []}
        },
        "global_risk_patterns": []
    },
    "it_services": {
        "name": "Типовая форма договора на IT-услуги",
        "code": "ТФ-ИТ-001",
        "version": "1.0",
        "date": "01.01.2025",
        "sections": {
            "1. ПРЕДМЕТ": {"required": True, "template": "IT-услуги.", "keywords": ["предмет", "it", "услуги"], "risk_patterns": []},
            "2. SLA": {"required": True, "template": "Уровень сервиса.", "keywords": ["sla", "уровень", "сервис"], "risk_patterns": [{"pattern": r"(?:нет|отсутств|без).*sla", "risk": "red", "issue": "Отсутствует SLA"}]},
            "3. СТОИМОСТЬ": {"required": True, "template": "Стоимость.", "keywords": ["стоимость", "оплата"], "risk_patterns": []},
            "4. ПРАВА НА РЕЗУЛЬТАТЫ": {"required": True, "template": "Права.", "keywords": ["права", "результат"], "risk_patterns": [{"pattern": r"прав\w+.*(?:принадлежат?|остают?ся).*исполнител", "risk": "red", "issue": "Права на ПО у Исполнителя"}]},
            "5. ЗАЩИТА ДАННЫХ": {"required": True, "template": "152-ФЗ.", "keywords": ["защита", "данные", "персональные"], "risk_patterns": []},
            "6. ПОДДЕРЖКА": {"required": True, "template": "Поддержка.", "keywords": ["поддержка"], "risk_patterns": []},
            "7. ОТВЕТСТВЕННОСТЬ": {"required": True, "template": "Ответственность.", "keywords": ["ответственность"], "risk_patterns": []},
            "8. СРОК И РАСТОРЖЕНИЕ": {"required": True, "template": "Срок.", "keywords": ["срок", "расторжение"], "risk_patterns": []},
            "9. РЕКВИЗИТЫ": {"required": True, "template": "Реквизиты.", "keywords": ["реквизиты", "подписи"], "risk_patterns": []}
        },
        "global_risk_patterns": []
    }
}

# ============================================================================
# ЗАГРУЗЧИК ДОКУМЕНТОВ
# ============================================================================

class DocumentLoader:
    """Загрузчик документов различных форматов"""
    
    @classmethod
    def load_txt(cls, file_content: bytes) -> str:
        """Загрузка TXT файла"""
        encodings = ['utf-8', 'cp1251', 'latin-1']
        for enc in encodings:
            try:
                return file_content.decode(enc)
            except:
                continue
        return file_content.decode('utf-8', errors='ignore')
    
    @classmethod
    def load_docx(cls, file_content: bytes) -> str:
        """Загрузка DOCX файла"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            return '\n'.join(text_parts)
        except ImportError:
            return "Ошибка: установите python-docx (pip install python-docx)"
        except Exception as e:
            return f"Ошибка чтения DOCX: {str(e)}"
    
    @classmethod
    def load_pdf(cls, file_content: bytes) -> str:
        """Загрузка PDF файла"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n'.join(text_parts)
        except ImportError:
            return "Ошибка: установите PyPDF2 (pip install PyPDF2)"
        except Exception as e:
            return f"Ошибка чтения PDF: {str(e)}"
    
    @classmethod
    def load_file(cls, uploaded_file) -> Tuple[bool, str]:
        """Универсальный загрузчик"""
        if not uploaded_file:
            return False, "Файл не выбран"
        
        try:
            content = uploaded_file.read()
            filename = uploaded_file.name.lower()
            
            if filename.endswith('.txt'):
                text = cls.load_txt(content)
            elif filename.endswith('.docx'):
                text = cls.load_docx(content)
            elif filename.endswith('.pdf'):
                text = cls.load_pdf(content)
            else:
                return False, "Неподдерживаемый формат"
            
            if text.startswith("Ошибка"):
                return False, text
            
            return True, text
        except Exception as e:
            return False, f"Ошибка: {str(e)}"

# ============================================================================
# АЛГОРИТМ СРАВНЕНИЯ
# ============================================================================

class AdvancedComparator:
    """Улучшенный алгоритм сравнения договоров"""
    
    STOP_WORDS = {'и', 'в', 'на', 'по', 'с', 'к', 'о', 'от', 'из', 'за', 'для', 'не', 'что', 
                  'как', 'это', 'все', 'или', 'при', 'до', 'без', 'его', 'её', 'их', 'быть',
                  'который', 'также', 'между', 'после', 'перед', 'через', 'более', 'менее',
                  'а', 'но', 'да', 'же', 'ли', 'бы', 'то', 'ни', 'если', 'чем'}
    
    @classmethod
    def tokenize(cls, text: str) -> List[str]:
        if not text:
            return []
        text = text.lower()
        text = re.sub(r'[^\w\sа-яё]', ' ', text)
        tokens = text.split()
        return [t for t in tokens if len(t) > 2 and t not in cls.STOP_WORDS]
    
    @classmethod
    def get_ngrams(cls, tokens: List[str], n: int = 2) -> Set[Tuple]:
        if len(tokens) < n:
            return set()
        return set(tuple(tokens[i:i+n]) for i in range(len(tokens) - n + 1))
    
    @classmethod
    def jaccard_similarity(cls, set1: Set, set2: Set) -> float:
        if not set1 or not set2:
            return 0.0
        intersection = len(set1 & set2)
        union = len(set1 | set2)
        return intersection / union if union > 0 else 0.0
    
    @classmethod
    def overlap_coefficient(cls, set1: Set, set2: Set) -> float:
        if not set1 or not set2:
            return 0.0
        intersection = len(set1 & set2)
        min_size = min(len(set1), len(set2))
        return intersection / min_size if min_size > 0 else 0.0
    
    @classmethod
    def levenshtein_ratio(cls, s1: str, s2: str) -> float:
        if not s1 or not s2:
            return 0.0
        return difflib.SequenceMatcher(None, s1.lower(), s2.lower()).ratio()
    
    @classmethod
    def compute_tf(cls, tokens: List[str]) -> Dict[str, float]:
        if not tokens:
            return {}
        counter = Counter(tokens)
        total = len(tokens)
        return {k: v / total for k, v in counter.items()}
    
    @classmethod
    def compute_idf(cls, documents: List[List[str]]) -> Dict[str, float]:
        if not documents:
            return {}
        n_docs = len(documents)
        df = Counter()
        for doc in documents:
            for token in set(doc):
                df[token] += 1
        return {token: math.log((n_docs + 1) / (count + 1)) + 1 for token, count in df.items()}
    
    @classmethod
    def cosine_similarity(cls, vec1: Dict[str, float], vec2: Dict[str, float]) -> float:
        if not vec1 or not vec2:
            return 0.0
        all_keys = set(vec1.keys()) | set(vec2.keys())
        dot_product = sum(vec1.get(k, 0) * vec2.get(k, 0) for k in all_keys)
        norm1 = math.sqrt(sum(v ** 2 for v in vec1.values()))
        norm2 = math.sqrt(sum(v ** 2 for v in vec2.values()))
        return dot_product / (norm1 * norm2) if norm1 > 0 and norm2 > 0 else 0.0
    
    @classmethod
    def keyword_match_score(cls, text: str, keywords: List[str]) -> Tuple[float, List[str]]:
        if not keywords:
            return 0.0, []
        text_lower = text.lower()
        matched = [kw for kw in keywords if kw.lower() in text_lower]
        return len(matched) / len(keywords), matched
    
    @classmethod
    def extract_sections(cls, text: str) -> Dict[str, str]:
        """Извлечение разделов из текста договора"""
        sections = {}
        patterns = [
            r'(\d+)\.\s*([А-ЯЁA-Z][А-ЯЁA-Z\s\-]+?)(?:\n|$)',
            r'(РАЗДЕЛ\s+\d+)[.\s]+(.+?)(?:\n|$)',
            r'(Статья\s+\d+)[.\s]+(.+?)(?:\n|$)',
        ]
        
        text_lines = text.split('\n')
        current_section = "ПРЕАМБУЛА"
        current_content = []
        
        for line in text_lines:
            found_section = False
            for pattern in patterns:
                match = re.match(pattern, line.strip(), re.IGNORECASE)
                if match:
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    current_section = f"{match.group(1)}. {match.group(2)}".strip().upper()
                    current_content = [line]
                    found_section = True
                    break
            
            if not found_section:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    @classmethod
    def compare_section(cls, contract_section: str, template_section: str, keywords: List[str]) -> Dict:
        result = {
            "found": False,
            "similarity_scores": {},
            "combined_score": 0,
            "matched_keywords": [],
            "keyword_score": 0,
            "deviations": [],
            "status": "missing"
        }
        
        if not contract_section:
            return result
        
        result["found"] = True
        
        contract_tokens = cls.tokenize(contract_section)
        template_tokens = cls.tokenize(template_section)
        
        contract_set = set(contract_tokens)
        template_set = set(template_tokens)
        jaccard = cls.jaccard_similarity(contract_set, template_set)
        result["similarity_scores"]["jaccard_words"] = jaccard
        
        contract_bigrams = cls.get_ngrams(contract_tokens, 2)
        template_bigrams = cls.get_ngrams(template_tokens, 2)
        jaccard_bigrams = cls.jaccard_similarity(contract_bigrams, template_bigrams)
        result["similarity_scores"]["jaccard_bigrams"] = jaccard_bigrams
        
        overlap = cls.overlap_coefficient(contract_set, template_set)
        result["similarity_scores"]["overlap"] = overlap
        
        levenshtein = cls.levenshtein_ratio(contract_section[:1000], template_section[:1000])
        result["similarity_scores"]["levenshtein"] = levenshtein
        
        docs = [contract_tokens, template_tokens]
        idf = cls.compute_idf(docs)
        contract_tfidf = {k: cls.compute_tf(contract_tokens).get(k, 0) * idf.get(k, 1) for k in contract_set}
        template_tfidf = {k: cls.compute_tf(template_tokens).get(k, 0) * idf.get(k, 1) for k in template_set}
        cosine = cls.cosine_similarity(contract_tfidf, template_tfidf)
        result["similarity_scores"]["cosine_tfidf"] = cosine
        
        keyword_score, matched = cls.keyword_match_score(contract_section, keywords)
        result["keyword_score"] = keyword_score
        result["matched_keywords"] = matched
        
        weights = {"jaccard_words": 0.15, "jaccard_bigrams": 0.20, "overlap": 0.15, 
                   "levenshtein": 0.20, "cosine_tfidf": 0.15, "keywords": 0.15}
        
        combined = (jaccard * weights["jaccard_words"] + jaccard_bigrams * weights["jaccard_bigrams"] +
                    overlap * weights["overlap"] + levenshtein * weights["levenshtein"] +
                    cosine * weights["cosine_tfidf"] + keyword_score * weights["keywords"])
        result["combined_score"] = combined
        
        if combined >= 0.7:
            result["status"] = "match"
        elif combined >= 0.4:
            result["status"] = "partial"
        else:
            result["status"] = "deviation"
        
        return result
    
    @classmethod
    def find_best_matching_section(cls, section_name: str, contract_sections: Dict[str, str], 
                                   template_text: str, keywords: List[str]) -> Tuple[str, Dict]:
        best_match = None
        best_result = None
        best_score = 0
        
        section_name_normalized = re.sub(r'^\d+\.\s*', '', section_name).lower()
        
        for contract_section_name, contract_section_text in contract_sections.items():
            contract_name_normalized = re.sub(r'^\d+\.\s*', '', contract_section_name).lower()
            name_similarity = cls.levenshtein_ratio(section_name_normalized, contract_name_normalized)
            
            comparison = cls.compare_section(contract_section_text, template_text, keywords)
            
            total_score = comparison["combined_score"] * 0.7 + name_similarity * 0.3
            
            if total_score > best_score:
                best_score = total_score
                best_match = contract_section_name
                best_result = comparison
                best_result["matched_section_name"] = contract_section_name
                best_result["name_similarity"] = name_similarity
        
        return best_match, best_result
    
    @classmethod
    def check_risk_patterns(cls, text: str, patterns: List[Dict]) -> List[Dict]:
        found_risks = []
        text_lower = text.lower()
        
        for p in patterns:
            match = re.search(p["pattern"], text_lower)
            if match:
                start = max(0, match.start() - 100)
                end = min(len(text_lower), match.end() + 100)
                context = text_lower[start:end]
                
                found_risks.append({
                    "issue": p["issue"],
                    "risk_level": p.get("risk", "yellow"),
                    "context": f"...{context}...",
                    "position": match.start()
                })
        
        return found_risks
    
    @classmethod
    def full_comparison(cls, contract_text: str, typical_form: Dict) -> Dict:
        result = {
            "form_name": typical_form.get("name", ""),
            "form_code": typical_form.get("code", ""),
            "form_version": typical_form.get("version", ""),
            "sections_analysis": [],
            "missing_sections": [],
            "found_sections": [],
            "partial_sections": [],
            "deviation_sections": [],
            "risks": [],
            "global_risks": [],
            "compliance_score": 0,
            "section_scores": {},
            "recommendations": [],
            "summary": ""
        }
        
        contract_sections = cls.extract_sections(contract_text)
        
        total_score = 0
        total_weight = 0
        
        for section_name, section_data in typical_form.get("sections", {}).items():
            weight = 2.0 if section_data.get("required", False) else 1.0
            total_weight += weight
            
            matched_name, comparison = cls.find_best_matching_section(
                section_name, contract_sections,
                section_data.get("template", ""),
                section_data.get("keywords", [])
            )
            
            section_result = {
                "section_name": section_name,
                "required": section_data.get("required", False),
                "matched_in_contract": matched_name,
                "comparison": comparison if comparison else {"found": False, "combined_score": 0, "status": "missing"},
                "risks": []
            }
            
            if matched_name and contract_sections.get(matched_name):
                section_risks = cls.check_risk_patterns(
                    contract_sections[matched_name],
                    section_data.get("risk_patterns", [])
                )
                section_result["risks"] = section_risks
                result["risks"].extend(section_risks)
            
            if comparison and comparison.get("found"):
                if comparison.get("status") == "match":
                    result["found_sections"].append(section_name)
                    total_score += weight
                elif comparison.get("status") == "partial":
                    result["partial_sections"].append(section_name)
                    total_score += weight * 0.6
                else:
                    result["deviation_sections"].append(section_name)
                    total_score += weight * 0.3
                    result["recommendations"].append(f"Раздел '{section_name}' существенно отличается от ТФ")
            else:
                if section_data.get("required"):
                    result["missing_sections"].append(section_name)
                    result["recommendations"].append(f"КРИТИЧНО: Отсутствует обязательный раздел '{section_name}'")
                else:
                    result["missing_sections"].append(section_name)
            
            result["sections_analysis"].append(section_result)
            if comparison:
                result["section_scores"][section_name] = comparison.get("combined_score", 0)
        
        global_risks = cls.check_risk_patterns(contract_text, typical_form.get("global_risk_patterns", []))
        result["global_risks"] = global_risks
        result["risks"].extend(global_risks)
        
        red_risks = sum(1 for r in result["risks"] if r.get("risk_level") == "red")
        yellow_risks = sum(1 for r in result["risks"] if r.get("risk_level") == "yellow")
        risk_penalty = red_risks * 0.15 + yellow_risks * 0.05
        
        if total_weight > 0:
            base_score = (total_score / total_weight) * 100
            result["compliance_score"] = max(0, base_score - risk_penalty * 100)
        
        if result["compliance_score"] >= 80:
            result["summary"] = "Высокое соответствие типовой форме."
        elif result["compliance_score"] >= 50:
            result["summary"] = "Частичное соответствие. Требуется проверка."
        else:
            result["summary"] = "Существенные отклонения. Требуется экспертиза."
        
        return result

# ============================================================================
# AI АНАЛИЗАТОР
# ============================================================================

class AIAnalyzer:
    @classmethod
    def generate_risk_analysis_prompt(cls, contract_text: str, comparison_result: Dict) -> str:
        risks_text = ""
        for risk in comparison_result.get("risks", []):
            level = "🔴 КРИТИЧЕСКИЙ" if risk.get("risk_level") == "red" else "🟡 ВНИМАНИЕ"
            risks_text += f"\n- {level}: {risk.get('issue', '')}\nКонтекст: {risk.get('context', '')[:200]}\n"
        
        missing = ", ".join(comparison_result.get("missing_sections", [])) or "Нет"
        deviations = ", ".join(comparison_result.get("deviation_sections", [])) or "Нет"
        
        prompt = f"""Ты — опытный юрист корпоративного права.

Проанализируй договор для АО «НПК» (Заказчик/Покупатель):

ДОГОВОР:
{contract_text[:8000]}

ОТКЛОНЕНИЯ ОТ ТИПОВОЙ ФОРМЫ:
- Соответствие ТФ: {comparison_result.get('compliance_score', 0):.0f}%
- Отсутствующие разделы: {missing}
- Разделы с отклонениями: {deviations}

ВЫЯВЛЕННЫЕ РИСКИ:
{risks_text if risks_text else "Автоматически не выявлены"}

Напиши юридическое заключение:

1. КРАТКОЕ РЕЗЮМЕ (2-3 предложения)
2. КРИТИЧЕСКИЕ РИСКИ (требуют устранения)
3. УСЛОВИЯ, ТРЕБУЮЩИЕ ВНИМАНИЯ
4. ОТСУТСТВУЮЩИЕ ЗАЩИТНЫЕ МЕХАНИЗМЫ
5. РЕКОМЕНДАЦИИ ПО ПЕРЕГОВОРАМ
6. ИТОГОВОЕ ЗАКЛЮЧЕНИЕ (согласовать/доработать/отклонить)
"""
        return prompt
    
    @classmethod
    def call_openai(cls, prompt: str, api_key: str) -> str:
        try:
            import requests
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": "gpt-4o-mini", "messages": [{"role": "user", "content": prompt}], "max_tokens": 4000},
                timeout=60
            )
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"]
            return f"Ошибка API: {response.status_code}"
        except Exception as e:
            return f"Ошибка: {str(e)}"
    
    @classmethod
    def call_anthropic(cls, prompt: str, api_key: str) -> str:
        try:
            import requests
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": api_key, "Content-Type": "application/json", "anthropic-version": "2023-06-01"},
                json={"model": "claude-3-haiku-20240307", "max_tokens": 4000, "messages": [{"role": "user", "content": prompt}]},
                timeout=60
            )
            if response.status_code == 200:
                return response.json()["content"][0]["text"]
            return f"Ошибка API: {response.status_code}"
        except Exception as e:
            return f"Ошибка: {str(e)}"

# ============================================================================
# ГЕНЕРАТОР ОТЧЕТОВ
# ============================================================================

class ReportGenerator:
    """Генератор отчетов в различных форматах"""
    
    @classmethod
    def generate_text_report(cls, data: Dict, user: Dict, ai_analysis: str = "") -> str:
        zone_names = {"green": "ЗЕЛЕНАЯ", "yellow": "ЖЕЛТАЯ", "red": "КРАСНАЯ"}
        
        report = f"""
================================================================================
                    ЗАКЛЮЧЕНИЕ ЮРИДИЧЕСКОГО ДЕПАРТАМЕНТА
                 АО «Новая перевозочная организация»
================================================================================

Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}
Исполнитель: {user.get('name', '')}
Должность: {user.get('position', '')}
Подразделение: {user.get('department', '')}
ID анализа: {data.get('analysis_id', '')}

--------------------------------------------------------------------------------
                         1. ОБЩИЕ СВЕДЕНИЯ О ДОГОВОРЕ
--------------------------------------------------------------------------------

Контрагент: {data.get('counterparty', 'Не указан')}
Номер договора: {data.get('contract_number', 'Не указан')}
Дата договора: {data.get('contract_date', 'Не указана')}
Тип договора: {data.get('contract_type', 'Не указан')}
Сумма договора: {data.get('amount', 0):,.2f} руб.
Статус согласования ЮД: {data.get('legal_status_label', 'Не определен')}

--------------------------------------------------------------------------------
                         2. РЕЗУЛЬТАТЫ АНАЛИЗА
--------------------------------------------------------------------------------

ЗОНА РИСКА: {zone_names.get(data.get('zone', ''), 'НЕ ОПРЕДЕЛЕНА')}
РИСК-СКОР: {data.get('risk_score', 0):.1f} / 10
СООТВЕТСТВИЕ ТИПОВОЙ ФОРМЕ: {data.get('compliance_score', 0):.1f}%

Требуется участие ЮД: {'ДА' if data.get('requires_legal') else 'НЕТ'}
Срок согласования: {data.get('deadline_days', 0)} рабочих дней
Согласующий: {data.get('approver', 'Не определен')}

--------------------------------------------------------------------------------
                   3. АНАЛИЗ ПО РАЗДЕЛАМ ТИПОВОЙ ФОРМЫ
--------------------------------------------------------------------------------

Типовая форма: {data.get('tf_name', 'Не выбрана')} ({data.get('tf_code', '')})

"""
        for s in data.get('sections_analysis', []):
            status = s.get('comparison', {}).get('status', 'missing')
            status_icon = {"match": "✅", "partial": "⚠️", "deviation": "❌", "missing": "❌"}.get(status, "❓")
            score = s.get('comparison', {}).get('combined_score', 0) * 100
            report += f"{status_icon} {s.get('section_name', '')} — Соответствие: {score:.0f}%\n"
            if s.get('risks'):
                for r in s['risks']:
                    lvl = "🔴" if r.get('risk_level') == 'red' else "🟡"
                    report += f"   {lvl} РИСК: {r.get('issue', '')}\n"
            report += "\n"
        
        if data.get('missing_sections'):
            report += "\n❌ ОТСУТСТВУЮЩИЕ РАЗДЕЛЫ:\n"
            for m in data['missing_sections']:
                report += f"   • {m}\n"
        
        report += """
--------------------------------------------------------------------------------
                         4. ВЫЯВЛЕННЫЕ РИСКИ
--------------------------------------------------------------------------------

"""
        red_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'red']
        if red_risks:
            report += "🔴 КРИТИЧЕСКИЕ РИСКИ:\n\n"
            for i, r in enumerate(red_risks, 1):
                report += f"   {i}. {r.get('issue', '')}\n"
                if r.get('context'):
                    report += f"      Контекст: {r.get('context', '')[:200]}...\n\n"
        else:
            report += "🔴 КРИТИЧЕСКИЕ РИСКИ: Не выявлены\n\n"
        
        yellow_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'yellow']
        if yellow_risks:
            report += "🟡 ТРЕБУЮТ ВНИМАНИЯ:\n\n"
            for i, r in enumerate(yellow_risks, 1):
                report += f"   {i}. {r.get('issue', '')}\n\n"
        else:
            report += "🟡 ТРЕБУЮТ ВНИМАНИЯ: Не выявлены\n\n"
        
        if ai_analysis:
            report += f"""
--------------------------------------------------------------------------------
                     5. ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ (AI)
--------------------------------------------------------------------------------

{ai_analysis}

"""
        
        if data.get('recommendations'):
            report += """
--------------------------------------------------------------------------------
                         6. РЕКОМЕНДАЦИИ
--------------------------------------------------------------------------------

"""
            for i, rec in enumerate(data['recommendations'], 1):
                report += f"   {i}. {rec}\n"
        
        report += f"""
--------------------------------------------------------------------------------
                         7. ИТОГОВОЕ ЗАКЛЮЧЕНИЕ
--------------------------------------------------------------------------------

{data.get('conclusion', data.get('summary', 'Заключение формируется на основе анализа.'))}

--------------------------------------------------------------------------------

Версия системы: Legal Traffic Light v5.1 Enterprise
Дата и время: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}

================================================================================
"""
        return report
    
    @classmethod
    def generate_docx_report(cls, data: Dict, user: Dict, ai_analysis: str = "") -> bytes:
        """Генерация DOCX отчета"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('ЗАКЛЮЧЕНИЕ ЮРИДИЧЕСКОГО ДЕПАРТАМЕНТА', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph('АО «Новая перевозочная организация»')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Метаданные
            doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph(f"Исполнитель: {user.get('name', '')}")
            doc.add_paragraph(f"Должность: {user.get('position', '')}")
            doc.add_paragraph(f"Подразделение: {user.get('department', '')}")
            doc.add_paragraph(f"ID анализа: {data.get('analysis_id', '')}")
            
            # Раздел 1
            doc.add_heading('1. ОБЩИЕ СВЕДЕНИЯ О ДОГОВОРЕ', level=1)
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            rows_data = [
                ("Контрагент", data.get('counterparty', 'Не указан')),
                ("Номер договора", data.get('contract_number', 'Не указан')),
                ("Дата договора", data.get('contract_date', 'Не указана')),
                ("Тип договора", data.get('contract_type', 'Не указан')),
                ("Сумма договора", f"{data.get('amount', 0):,.2f} руб."),
                ("Статус согласования", data.get('legal_status_label', 'Не определен')),
            ]
            
            for i, (label, value) in enumerate(rows_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            # Раздел 2
            doc.add_heading('2. РЕЗУЛЬТАТЫ АНАЛИЗА', level=1)
            
            zone_names = {"green": "🟢 ЗЕЛЕНАЯ", "yellow": "🟡 ЖЕЛТАЯ", "red": "🔴 КРАСНАЯ"}
            zone = data.get('zone', '')
            
            p = doc.add_paragraph()
            p.add_run(f"ЗОНА РИСКА: {zone_names.get(zone, 'НЕ ОПРЕДЕЛЕНА')}").bold = True
            
            doc.add_paragraph(f"Риск-скор: {data.get('risk_score', 0):.1f} / 10")
            doc.add_paragraph(f"Соответствие ТФ: {data.get('compliance_score', 0):.1f}%")
            doc.add_paragraph(f"Требуется участие ЮД: {'ДА' if data.get('requires_legal') else 'НЕТ'}")
            doc.add_paragraph(f"Срок согласования: {data.get('deadline_days', 0)} рабочих дней")
            
            # Раздел 3
            doc.add_heading('3. АНАЛИЗ ПО РАЗДЕЛАМ', level=1)
            doc.add_paragraph(f"Типовая форма: {data.get('tf_name', '')} ({data.get('tf_code', '')})")
            
            for s in data.get('sections_analysis', []):
                status = s.get('comparison', {}).get('status', 'missing')
                status_icon = {"match": "✅", "partial": "⚠️", "deviation": "❌", "missing": "❌"}.get(status, "❓")
                score = s.get('comparison', {}).get('combined_score', 0) * 100
                
                p = doc.add_paragraph()
                p.add_run(f"{status_icon} {s.get('section_name', '')}").bold = True
                p.add_run(f" — {score:.0f}%")
                
                for r in s.get('risks', []):
                    lvl = "🔴" if r.get('risk_level') == 'red' else "🟡"
                    doc.add_paragraph(f"   {lvl} {r.get('issue', '')}", style='List Bullet')
            
            # Раздел 4
            doc.add_heading('4. ВЫЯВЛЕННЫЕ РИСКИ', level=1)
            
            red_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'red']
            if red_risks:
                doc.add_paragraph("КРИТИЧЕСКИЕ РИСКИ:", style='Intense Quote')
                for r in red_risks:
                    doc.add_paragraph(f"• {r.get('issue', '')}")
            
            yellow_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'yellow']
            if yellow_risks:
                doc.add_paragraph("ТРЕБУЮТ ВНИМАНИЯ:", style='Intense Quote')
                for r in yellow_risks:
                    doc.add_paragraph(f"• {r.get('issue', '')}")
            
            # AI анализ
            if ai_analysis:
                doc.add_heading('5. ЭКСПЕРТНОЕ ЗАКЛЮЧЕНИЕ (AI)', level=1)
                doc.add_paragraph(ai_analysis)
            
            # Рекомендации
            if data.get('recommendations'):
                doc.add_heading('6. РЕКОМЕНДАЦИИ', level=1)
                for rec in data['recommendations']:
                    doc.add_paragraph(rec, style='List Number')
            
            # Заключение
            doc.add_heading('7. ИТОГОВОЕ ЗАКЛЮЧЕНИЕ', level=1)
            doc.add_paragraph(data.get('conclusion', data.get('summary', '')))
            
            # Футер
            doc.add_paragraph()
            doc.add_paragraph(f"Legal Traffic Light v5.1 | {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            
            # Сохраняем в байты
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            return b"Error: Install python-docx"
        except Exception as e:
            return f"Error: {str(e)}".encode()
    
    @classmethod
    def generate_pdf_report(cls, data: Dict, user: Dict, ai_analysis: str = "") -> bytes:
        """Генерация PDF отчета"""
        try:
            from fpdf import FPDF
            
            class PDF(FPDF):
                def header(self):
                    self.set_font('DejaVu', 'B', 12)
                    self.cell(0, 10, 'ЗАКЛЮЧЕНИЕ ЮРИДИЧЕСКОГО ДЕПАРТАМЕНТА', 0, 1, 'C')
                    self.set_font('DejaVu', '', 10)
                    self.cell(0, 5, 'АО «Новая перевозочная организация»', 0, 1, 'C')
                    self.ln(5)
                
                def footer(self):
                    self.set_y(-15)
                    self.set_font('DejaVu', 'I', 8)
                    self.cell(0, 10, f'Legal Traffic Light v5.1 | Страница {self.page_no()}', 0, 0, 'C')
            
            pdf = PDF()
            
            # Добавляем шрифт с поддержкой кириллицы
            # Используем встроенный шрифт
            pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
            pdf.add_font('DejaVu', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', uni=True)
            pdf.add_font('DejaVu', 'I', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf', uni=True)
            
            pdf.add_page()
            pdf.set_font('DejaVu', '', 10)
            
            # Метаданные
            pdf.cell(0, 6, f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}", 0, 1)
            pdf.cell(0, 6, f"Исполнитель: {user.get('name', '')}", 0, 1)
            pdf.cell(0, 6, f"ID анализа: {data.get('analysis_id', '')}", 0, 1)
            pdf.ln(5)
            
            # Раздел 1
            pdf.set_font('DejaVu', 'B', 11)
            pdf.cell(0, 8, '1. ОБЩИЕ СВЕДЕНИЯ О ДОГОВОРЕ', 0, 1)
            pdf.set_font('DejaVu', '', 10)
            
            pdf.cell(0, 6, f"Контрагент: {data.get('counterparty', 'Не указан')}", 0, 1)
            pdf.cell(0, 6, f"Сумма: {data.get('amount', 0):,.2f} руб.", 0, 1)
            pdf.cell(0, 6, f"Тип: {data.get('contract_type', '')}", 0, 1)
            pdf.ln(5)
            
            # Раздел 2
            pdf.set_font('DejaVu', 'B', 11)
            pdf.cell(0, 8, '2. РЕЗУЛЬТАТЫ АНАЛИЗА', 0, 1)
            pdf.set_font('DejaVu', '', 10)
            
            zone_names = {"green": "ЗЕЛЕНАЯ", "yellow": "ЖЕЛТАЯ", "red": "КРАСНАЯ"}
            pdf.cell(0, 6, f"Зона риска: {zone_names.get(data.get('zone', ''), 'НЕ ОПРЕДЕЛЕНА')}", 0, 1)
            pdf.cell(0, 6, f"Риск-скор: {data.get('risk_score', 0):.1f} / 10", 0, 1)
            pdf.cell(0, 6, f"Соответствие ТФ: {data.get('compliance_score', 0):.1f}%", 0, 1)
            pdf.ln(5)
            
            # Раздел 3 - Риски
            pdf.set_font('DejaVu', 'B', 11)
            pdf.cell(0, 8, '3. ВЫЯВЛЕННЫЕ РИСКИ', 0, 1)
            pdf.set_font('DejaVu', '', 10)
            
            red_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'red']
            if red_risks:
                pdf.set_font('DejaVu', 'B', 10)
                pdf.cell(0, 6, 'КРИТИЧЕСКИЕ:', 0, 1)
                pdf.set_font('DejaVu', '', 9)
                for r in red_risks:
                    pdf.multi_cell(0, 5, f"• {r.get('issue', '')}")
            
            yellow_risks = [r for r in data.get('risks', []) if r.get('risk_level') == 'yellow']
            if yellow_risks:
                pdf.set_font('DejaVu', 'B', 10)
                pdf.cell(0, 6, 'ВНИМАНИЕ:', 0, 1)
                pdf.set_font('DejaVu', '', 9)
                for r in yellow_risks:
                    pdf.multi_cell(0, 5, f"• {r.get('issue', '')}")
            
            if not red_risks and not yellow_risks:
                pdf.cell(0, 6, 'Риски не выявлены', 0, 1)
            
            pdf.ln(5)
            
            # Заключение
            pdf.set_font('DejaVu', 'B', 11)
            pdf.cell(0, 8, '4. ЗАКЛЮЧЕНИЕ', 0, 1)
            pdf.set_font('DejaVu', '', 10)
            pdf.multi_cell(0, 6, data.get('conclusion', data.get('summary', '')))
            
            return pdf.output(dest='S').encode('latin-1')
            
        except ImportError:
            # Fallback - возвращаем текстовый отчет
            return cls.generate_text_report(data, user, ai_analysis).encode('utf-8')
        except Exception as e:
            return f"Error generating PDF: {str(e)}".encode()
    
    @classmethod
    def generate_json_for_1c(cls, data: Dict, user: Dict) -> str:
        """JSON для 1С"""
        report = {
            "Документ": "ЗаключениеЮД",
            "Версия": "2.0",
            "ДатаСоздания": datetime.now().isoformat(),
            "Организация": "АО НПК",
            "Автор": {"ФИО": user.get("name", ""), "Должность": user.get("position", ""), "Подразделение": user.get("department", "")},
            "ПараметрыДоговора": {
                "Номер": data.get("contract_number", ""),
                "Дата": data.get("contract_date", ""),
                "Контрагент": data.get("counterparty", ""),
                "Сумма": data.get("amount", 0),
                "Валюта": "RUB",
                "ТипДоговора": data.get("contract_type", ""),
            },
            "РезультатАнализа": {
                "ЗонаРиска": data.get("zone", ""),
                "РискСкор": data.get("risk_score", 0),
                "СоответствиеТФ": data.get("compliance_score", 0),
                "КодТФ": data.get("tf_code", ""),
                "СтатусСогласованияЮД": data.get("legal_status", ""),
                "СрокСогласования": data.get("deadline_days", 0),
                "Согласующий": data.get("approver", ""),
                "ТребуетсяУчастиеЮД": data.get("requires_legal", False)
            },
            "АнализРазделов": [
                {
                    "Раздел": s.get("section_name", ""),
                    "Обязательный": s.get("required", False),
                    "СоответствиеПроцент": s.get("comparison", {}).get("combined_score", 0) * 100,
                    "Статус": s.get("comparison", {}).get("status", "missing"),
                    "Риски": [{"Уровень": r.get("risk_level", ""), "Описание": r.get("issue", "")} for r in s.get("risks", [])]
                }
                for s in data.get("sections_analysis", [])
            ],
            "ВыявленныеРиски": {
                "Критические": [{"Описание": r.get("issue", ""), "Контекст": r.get("context", "")[:500]} for r in data.get("risks", []) if r.get("risk_level") == "red"],
                "Предупреждения": [{"Описание": r.get("issue", ""), "Контекст": r.get("context", "")[:500]} for r in data.get("risks", []) if r.get("risk_level") == "yellow"]
            },
            "ОтсутствующиеРазделы": data.get("missing_sections", []),
            "Рекомендации": data.get("recommendations", []),
            "Заключение": data.get("conclusion", ""),
            "СлужебнаяИнформация": {"ИдентификаторАнализа": data.get("analysis_id", ""), "ВерсияСистемы": "5.1"}
        }
        return json.dumps(report, ensure_ascii=False, indent=2)
    
    @classmethod
    def generate_json_knowledge_base(cls, data: Dict, user: Dict, contract_text: str = "") -> str:
        """JSON для базы знаний"""
        kb_entry = {
            "id": data.get("analysis_id", ""),
            "timestamp": datetime.now().isoformat(),
            "type": "contract_analysis",
            "metadata": {
                "counterparty": data.get("counterparty", ""),
                "contract_number": data.get("contract_number", ""),
                "contract_date": data.get("contract_date", ""),
                "contract_type": data.get("contract_type", ""),
                "amount": data.get("amount", 0),
                "analyst": user.get("name", ""),
                "department": user.get("department", ""),
            },
            "analysis": {
                "risk_zone": data.get("zone", ""),
                "risk_score": data.get("risk_score", 0),
                "compliance_score": data.get("compliance_score", 0),
                "typical_form": {"name": data.get("tf_name", ""), "code": data.get("tf_code", "")},
                "legal_status": data.get("legal_status", ""),
                "requires_legal_review": data.get("requires_legal", False),
            },
            "risks": {
                "critical": [{"issue": r.get("issue", ""), "context": r.get("context", "")[:300]} for r in data.get("risks", []) if r.get("risk_level") == "red"],
                "warnings": [{"issue": r.get("issue", ""), "context": r.get("context", "")[:300]} for r in data.get("risks", []) if r.get("risk_level") == "yellow"]
            },
            "sections": {
                "missing": data.get("missing_sections", []),
                "deviations": data.get("deviation_sections", []),
                "analysis": [
                    {"name": s.get("section_name", ""), "score": s.get("comparison", {}).get("combined_score", 0), "status": s.get("comparison", {}).get("status", ""), "risks": len(s.get("risks", []))}
                    for s in data.get("sections_analysis", [])
                ]
            },
            "recommendations": data.get("recommendations", []),
            "conclusion": data.get("conclusion", ""),
            "contract_text_hash": hashlib.md5(contract_text.encode()).hexdigest() if contract_text else "",
            "version": "5.1"
        }
        return json.dumps(kb_entry, ensure_ascii=False, indent=2)

# ============================================================================
# БЕЗОПАСНОСТЬ
# ============================================================================

class SecurityValidator:
    ALLOWED_EXTENSIONS = {'.txt', '.docx', '.pdf', '.doc', '.json'}
    MAX_FILE_SIZE = 15 * 1024 * 1024
    MAX_TEXT_LENGTH = 1_000_000
    
    DANGEROUS_PATTERNS = [r'<script[^>]*>.*?</script>', r'javascript:', r'on\w+\s*=', r'<iframe[^>]*>', r'eval\s*\(']
    
    @classmethod
    def sanitize_text(cls, text: str) -> str:
        if not text:
            return ""
        text = html.escape(text)
        for pattern in cls.DANGEROUS_PATTERNS:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)
        return text[:cls.MAX_TEXT_LENGTH]
    
    @classmethod
    def validate_file(cls, f) -> Tuple[bool, str]:
        if not f:
            return False, "Файл не выбран"
        if f.size > cls.MAX_FILE_SIZE:
            return False, f"Файл > {cls.MAX_FILE_SIZE // 1024 // 1024} MB"
        ext = '.' + f.name.split('.')[-1].lower()
        if ext not in cls.ALLOWED_EXTENSIONS:
            return False, f"Разрешены: {', '.join(cls.ALLOWED_EXTENSIONS)}"
        return True, "OK"
    
    @classmethod
    def validate_user(cls, name: str, position: str, department: str) -> Tuple[bool, str]:
        if not name or len(name.strip()) < 5:
            return False, "ФИО: минимум 5 символов"
        if not position:
            return False, "Выберите должность"
        if not department:
            return False, "Выберите подразделение"
        return True, "OK"
    
    @classmethod
    def validate_amount(cls, value: str) -> Tuple[bool, float, str]:
        try:
            cleaned = re.sub(r'[^\d.,]', '', value).replace(',', '.')
            if not cleaned:
                return True, 0, "OK"
            amount = float(cleaned)
            return True, amount, "OK"
        except:
            return False, 0, "Некорректный формат"

# ============================================================================
# ОПРЕДЕЛЕНИЕ ЗОНЫ РИСКА
# ============================================================================

@dataclass
class AnalysisInput:
    amount: float = 0
    document_form: DocumentForm = DocumentForm.TYPICAL
    document_type: str = ""
    deal_type: str = ""
    legal_status: LegalStatus = LegalStatus.NOT_SUBMITTED
    is_single_supplier: bool = False
    is_tender: bool = False
    tender_amount: float = 0
    contract_years: int = 0
    changes_essential: bool = False
    is_urgent: bool = False
    counterparty: str = ""
    contract_number: str = ""
    contract_date: str = ""

@dataclass
class ZoneResult:
    zone: RiskZone = RiskZone.GREEN
    zone_reason: str = ""
    deadline_days: int = 0
    requires_legal: bool = False
    responsible: str = "Инициатор"
    recommendations: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    red_flags: List[str] = field(default_factory=list)
    yellow_flags: List[str] = field(default_factory=list)
    green_flags: List[str] = field(default_factory=list)
    required_documents: List[str] = field(default_factory=list)
    approval_route: List[str] = field(default_factory=list)

def determine_risk_zone(inp: AnalysisInput) -> ZoneResult:
    result = ZoneResult()
    
    # КРАСНАЯ по типу сделки
    if inp.deal_type and inp.deal_type in RED_ZONE_ALWAYS:
        result.zone = RiskZone.RED
        result.zone_reason = f"'{inp.deal_type}' — всегда красная зона (п. 4.3)"
        result.red_flags.append(f"Стратегическая сделка: {inp.deal_type}")
        result.requires_legal = True
        result.responsible = "Юридический департамент"
        result.deadline_days = Deadlines.EXTENDED
        result.approval_route = ["ЮД (этап планирования)", "Совместная работа"]
        result.required_documents = ["Описание задачи", "Проект договора", "ТЗ", "КП", "Карточка контрагента"]
        return result
    
    if inp.document_type and inp.document_type in RED_DOCUMENTS_ALWAYS:
        result.zone = RiskZone.RED
        result.zone_reason = f"'{inp.document_type}' — всегда красная зона"
        result.red_flags.append(f"Критический документ: {inp.document_type}")
        result.requires_legal = True
        result.responsible = "Юридический департамент"
        result.deadline_days = Deadlines.URGENT
        result.approval_route = ["ЮД НЕЗАМЕДЛИТЕЛЬНО"]
        result.warnings.append("⚠️ НЕМЕДЛЕННОЕ привлечение ЮД!")
        return result
    
    if inp.is_tender and inp.tender_amount > Thresholds.TENDER_RED:
        result.zone = RiskZone.RED
        result.zone_reason = f"Тендер > 3 000 000 руб."
        result.red_flags.append(f"Тендер {inp.tender_amount:,.0f} руб.")
        result.requires_legal = True
        result.responsible = "Юридический департамент"
        result.deadline_days = Deadlines.EXTENDED
        return result
    
    if inp.amount > Thresholds.YELLOW_MAX:
        result.zone = RiskZone.RED
        result.zone_reason = f"Сумма {inp.amount:,.0f} руб. > 5 000 000"
        result.red_flags.append("Сумма > 5 млн руб.")
        result.requires_legal = True
        result.responsible = "Юридический департамент"
        result.deadline_days = Deadlines.EXTENDED
        result.approval_route = ["ЮД (планирование)", "Совместная работа"]
        return result
    
    # ЖЕЛТАЯ зона
    yellow_reasons = []
    
    if inp.deal_type and inp.deal_type in YELLOW_ZONE_TYPES:
        yellow_reasons.append(inp.deal_type)
        result.yellow_flags.append(inp.deal_type)
    
    if inp.is_single_supplier and inp.amount > Thresholds.SINGLE_SUPPLIER_YELLOW:
        yellow_reasons.append("Единственный поставщик > 100К")
        result.yellow_flags.append("Единственный поставщик")
    
    if inp.contract_years > Thresholds.CONTRACT_CONTROL_YEARS:
        yellow_reasons.append(f"Контрольный срок > {Thresholds.CONTRACT_CONTROL_YEARS} лет")
        result.yellow_flags.append(f"Договор > {inp.contract_years} лет")
    
    if inp.document_form == DocumentForm.TYPICAL:
        if inp.amount > Thresholds.GREEN_TF_MAX:
            yellow_reasons.append(f"ТФ > 100К")
            result.yellow_flags.append(f"ТФ: {inp.amount:,.0f} руб.")
        else:
            result.green_flags.append("ТФ ≤ 100К")
    elif inp.document_form in [DocumentForm.COUNTERPARTY, DocumentForm.FREE, DocumentForm.SELF_DEVELOPED]:
        if inp.amount > Thresholds.GREEN_NON_TF_MAX:
            yellow_reasons.append("Нетиповая форма > 50К")
            result.yellow_flags.append("Нетиповая > 50К")
        else:
            result.green_flags.append("Нетиповая ≤ 50К")
    elif inp.document_form == DocumentForm.MODIFIED_TF:
        if inp.changes_essential:
            yellow_reasons.append("Изменение существенных условий")
            result.yellow_flags.append("Существенные изменения ТФ")
        elif inp.amount > Thresholds.GREEN_NON_TF_MAX:
            yellow_reasons.append("Изменения ТФ > 50К")
            result.yellow_flags.append("Изменения ТФ > 50К")
        else:
            result.green_flags.append("Изменения ТФ ≤ 50К")
    
    if yellow_reasons:
        result.zone = RiskZone.YELLOW
        result.zone_reason = "; ".join(yellow_reasons)
        result.requires_legal = True
        result.responsible = "ЮД (экспертиза)"
        result.deadline_days = Deadlines.STANDARD
        result.approval_route = ["Инициатор → СЭД → ЮД (5 дн.) → Подписание"]
        result.required_documents = ["Проект договора", "ТЗ", "КП", "Карточка контрагента"]
        result.warnings.append("⚠️ ЗАПРЕЩЕНО подписание до визы ЮД!")
        return result
    
    # ЗЕЛЕНАЯ
    result.zone = RiskZone.GREEN
    result.zone_reason = "Зеленый коридор (п. 4.1)"
    result.requires_legal = False
    result.responsible = "Инициатор"
    result.approval_route = ["Инициатор → Руководитель → Подписание"]
    result.recommendations.append("Используйте актуальную ТФ без изменений")
    
    return result

# ============================================================================
# ИСТОРИЯ
# ============================================================================

def add_to_history(data: Dict):
    if "history" not in st.session_state:
        st.session_state.history = []
    data["id"] = hashlib.md5(f"{datetime.now().isoformat()}{data.get('counterparty', '')}".encode()).hexdigest()[:8]
    data["timestamp"] = datetime.now().isoformat()
    st.session_state.history.insert(0, data)
    if len(st.session_state.history) > 100:
        st.session_state.history = st.session_state.history[:100]

# ============================================================================
# СТИЛИ
# ============================================================================

def apply_styles():
    st.markdown("""
    <style>
        :root { --red: #dc3545; --yellow: #ffc107; --green: #28a745; --blue: #0d6efd; --purple: #6f42c1; }
        .main-header { background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%); padding: 1.5rem 2rem; border-radius: 12px; margin-bottom: 1.5rem; color: white; }
        .traffic-light { display: flex; gap: 8px; margin-bottom: 0.5rem; }
        .traffic-light span { width: 18px; height: 18px; border-radius: 50%; }
        .tl-red { background: var(--red); box-shadow: 0 0 10px var(--red); }
        .tl-yellow { background: var(--yellow); box-shadow: 0 0 10px var(--yellow); }
        .tl-green { background: var(--green); box-shadow: 0 0 10px var(--green); }
        .version-badge { background: var(--purple); padding: 2px 10px; border-radius: 10px; font-size: 0.7rem; margin-left: 10px; }
        .zone-card { border-radius: 12px; padding: 1.5rem; margin-bottom: 1rem; border-left: 5px solid; }
        .zone-card.green { background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%); border-left-color: var(--green); }
        .zone-card.yellow { background: linear-gradient(135deg, #fff3cd 0%, #ffeeba 100%); border-left-color: var(--yellow); }
        .zone-card.red { background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%); border-left-color: var(--red); }
        .zone-title { font-size: 1.2rem; font-weight: bold; margin-bottom: 0.5rem; }
        .risk-item { background: #f8f9fa; border-radius: 8px; padding: 1rem; margin-bottom: 0.8rem; border-left: 3px solid; }
        .risk-item.red { border-left-color: var(--red); background: #fff5f5; }
        .risk-item.yellow { border-left-color: var(--yellow); background: #fffcf0; }
        .risk-item.green { border-left-color: var(--green); background: #f0fff4; }
        .section-score { display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: bold; }
        .score-high { background: #d4edda; color: #155724; }
        .score-medium { background: #fff3cd; color: #856404; }
        .score-low { background: #f8d7da; color: #721c24; }
        footer { visibility: hidden; }
    </style>
    """, unsafe_allow_html=True)

# ============================================================================
# SESSION STATE
# ============================================================================

def init_session():
    defaults = {
        "authenticated": False, "user": None, "demo_mode": False,
        "contract_text": "", "zone_result": None, "comparison_result": None,
        "ai_analysis": "", "history": [], "current_input": None,
        "custom_typical_forms": {}, "selected_tf": None
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

# ============================================================================
# ДЕМО
# ============================================================================

DEMO_CONTRACT = """ДОГОВОР ОКАЗАНИЯ УСЛУГ № 2025/IT-001

г. Москва                                                     «10» января 2025 г.

ООО «ИТ-Решения» и АО «Новая перевозочная организация» заключили договор:

1. ПРЕДМЕТ ДОГОВОРА
1.1. Исполнитель обязуется оказать услуги по технической поддержке ИС.

2. СТОИМОСТЬ И ПОРЯДОК РАСЧЕТОВ
2.1. Стоимость: 6 000 000 рублей.
2.2. Оплата в течение 5 дней после подписания Акта.
2.3. Неустойка за просрочку оплаты: 0,5% в день.

3. СРОКИ ОКАЗАНИЯ УСЛУГ
3.1. С 01.01.2025 по 31.12.2025.

4. ПОРЯДОК СДАЧИ-ПРИЕМКИ
4.1. Акт подписывается в течение 2 рабочих дней.

5. ПРАВА И ОБЯЗАННОСТИ
5.1. Исполнитель оказывает услуги. Заказчик оплачивает.

6. ОТВЕТСТВЕННОСТЬ СТОРОН
6.1. Максимальная ответственность Исполнителя ограничена суммой последнего месячного платежа.
6.2. Исполнитель не несет ответственности за косвенные убытки и упущенную выгоду.

7. КОНФИДЕНЦИАЛЬНОСТЬ
7.1. Срок: 5 лет после прекращения.
7.2. Штраф за нарушение: 10 000 000 рублей.

8. ИНТЕЛЛЕКТУАЛЬНАЯ СОБСТВЕННОСТЬ
8.1. Все результаты работ являются исключительной собственностью Исполнителя.
8.2. Заказчику предоставляется неисключительная лицензия.

9. СРОК ДЕЙСТВИЯ И РАСТОРЖЕНИЕ
9.1. До 31.12.2025.
9.2. Исполнитель вправе расторгнуть договор в одностороннем порядке за 5 дней.
9.3. Заказчик вправе расторгнуть только при существенном нарушении.

10. РАЗРЕШЕНИЕ СПОРОВ
10.1. Арбитражный суд г. Санкт-Петербурга.

11. ФОРС-МАЖОР
11.1. Непреодолимая сила освобождает от ответственности.

12. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ
12.1. Субподряд без согласия Заказчика.
12.2. Исполнитель может изменить цену в одностороннем порядке за 10 дней.

13. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН
ЗАКАЗЧИК: АО «НПК»...
ИСПОЛНИТЕЛЬ: ООО «ИТ-Решения»..."""

DEMO_USER = {"name": "Демо-пользователь", "position": "Специалист", "department": "Юридический департамент"}

# ============================================================================
# UI КОМПОНЕНТЫ
# ============================================================================

def render_auth():
    st.markdown("""
    <div style="text-align: center; padding: 2rem;">
        <div class="traffic-light" style="justify-content: center;">
            <span class="tl-red"></span><span class="tl-yellow"></span><span class="tl-green"></span>
        </div>
        <h1>🚦 Legal Traffic Light</h1>
        <p style="color: #6c757d;">АО «Новая перевозочная организация»</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("### 🔐 Вход")
        with st.form("auth"):
            name = st.text_input("ФИО *", placeholder="Иванов Иван Иванович")
            position = st.selectbox("Должность *", [""] + POSITIONS)
            department = st.selectbox("Подразделение *", [""] + DEPARTMENTS)
            col_a, col_b = st.columns(2)
            with col_a:
                if st.form_submit_button("🚀 Войти", use_container_width=True, type="primary"):
                    valid, msg = SecurityValidator.validate_user(name, position, department)
                    if valid:
                        st.session_state.authenticated = True
                        st.session_state.user = {"name": name, "position": position, "department": department}
                        st.rerun()
                    else:
                        st.error(msg)
            with col_b:
                if st.form_submit_button("📋 Демо", use_container_width=True):
                    st.session_state.authenticated = True
                    st.session_state.user = DEMO_USER.copy()
                    st.session_state.demo_mode = True
                    st.rerun()

def render_sidebar():
    with st.sidebar:
        st.markdown("### 👤 Пользователь")
        if st.session_state.demo_mode:
            st.info("🎯 Демо-режим")
        u = st.session_state.user
        st.markdown(f"**{u['name']}**\n\n{u['position']}\n\n_{u['department']}_")
        if st.button("🚪 Выйти", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()
        
        st.markdown("---")
        st.markdown("### 📋 Матрица")
        st.markdown("🟢 ТФ ≤100К / Нетип. ≤50К\n🟡 ≤5М / Особые\n🔴 >5М / Стратегические")
        
        if st.session_state.history:
            st.markdown("---")
            st.markdown("### 📜 История")
            for h in st.session_state.history[:5]:
                z = {"green": "🟢", "yellow": "🟡", "red": "🔴"}.get(h.get("zone"), "⚪")
                st.markdown(f"{z} {h.get('counterparty', 'N/A')[:15]} | {h.get('timestamp', '')[:10]}")

def render_analysis_tab():
    st.markdown("### 📝 Параметры договора")
    
    col1, col2 = st.columns(2)
    with col1:
        counterparty = st.text_input("Контрагент", placeholder="ООО «Название»")
        contract_number = st.text_input("Номер договора", placeholder="2025/001")
        contract_date = st.date_input("Дата", value=date.today())
        doc_type = st.selectbox("Тип договора", ["Договор оказания услуг", "Договор поставки", "Договор аренды", "Договор аренды вагонов", "IT-услуги", "Другой"])
        doc_form = st.selectbox("Форма", [
            ("ТФ без изменений", DocumentForm.TYPICAL),
            ("Форма контрагента", DocumentForm.COUNTERPARTY),
            ("Свободная форма", DocumentForm.FREE),
            ("ТФ с изменениями", DocumentForm.MODIFIED_TF),
            ("Самостоятельная разработка", DocumentForm.SELF_DEVELOPED),
        ], format_func=lambda x: x[0])[1]
    
    with col2:
        amount_str = st.text_input("Сумма (₽)", placeholder="1500000")
        valid_amount, amount, _ = SecurityValidator.validate_amount(amount_str) if amount_str else (True, 0, "")
        
        legal_status = st.selectbox("Статус согласования ЮД", list(LegalStatus), format_func=lambda x: LEGAL_STATUS_LABELS[x])
        deal_type = st.selectbox("Специальный тип", ["Не применимо"] + RED_ZONE_ALWAYS + YELLOW_ZONE_TYPES)
        if deal_type == "Не применимо":
            deal_type = ""
        
        is_single_supplier = st.checkbox("Единственный поставщик")
        is_tender = st.checkbox("Тендер")
        tender_amount = 0
        if is_tender:
            t_str = st.text_input("Сумма тендера", "")
            if t_str:
                _, tender_amount, _ = SecurityValidator.validate_amount(t_str)
    
    with st.expander("🔧 Дополнительно"):
        col_a, col_b = st.columns(2)
        with col_a:
            contract_years = st.number_input("Срок (лет)", 0, 50, 0)
            changes_essential = st.checkbox("Существенные изменения")
        with col_b:
            is_urgent = st.checkbox("🚨 Срочно")
            if is_urgent:
                st.warning("Требуется обоснование!")
    
    st.markdown("---")
    
    col_b1, col_b2, col_b3 = st.columns(3)
    with col_b1:
        if st.button("🚦 Определить зону", type="primary", use_container_width=True):
            inp = AnalysisInput(
                amount=amount, document_form=doc_form, document_type=doc_type, deal_type=deal_type,
                legal_status=legal_status, is_single_supplier=is_single_supplier, is_tender=is_tender,
                tender_amount=tender_amount, contract_years=contract_years, changes_essential=changes_essential,
                is_urgent=is_urgent, counterparty=counterparty, contract_number=contract_number, contract_date=str(contract_date)
            )
            st.session_state.zone_result = determine_risk_zone(inp)
            st.session_state.current_input = inp
            st.rerun()
    with col_b2:
        if st.button("📝 Демо-договор", use_container_width=True):
            st.session_state.contract_text = DEMO_CONTRACT
            st.rerun()
    with col_b3:
        if st.button("🗑️ Очистить", use_container_width=True):
            st.session_state.zone_result = None
            st.session_state.comparison_result = None
            st.session_state.ai_analysis = ""
            st.session_state.contract_text = ""
            st.rerun()
    
    # Отображение результата зоны
    if st.session_state.zone_result:
        zr = st.session_state.zone_result
        zone_cfg = {
            RiskZone.GREEN: ("green", "🟢 ЗЕЛЕНАЯ", "Самостоятельно"),
            RiskZone.YELLOW: ("yellow", "🟡 ЖЕЛТАЯ", "Согласование ЮД"),
            RiskZone.RED: ("red", "🔴 КРАСНАЯ", "Полное сопровождение"),
        }
        
        # Безопасное получение конфигурации зоны
        if zr.zone in zone_cfg:
            zc, zt, zs = zone_cfg[zr.zone]
        else:
            zc, zt, zs = "green", "⚪ НЕ ОПРЕДЕЛЕНА", "Проверьте параметры"
        
        st.markdown(f'<div class="zone-card {zc}"><div class="zone-title">{zt}</div><div>{zs}</div><div style="font-size: 0.9rem; color: #555;">{zr.zone_reason}</div></div>', unsafe_allow_html=True)
        st.markdown(f"**Статус ЮД:** {LEGAL_STATUS_LABELS[legal_status]}")
        
        cols = st.columns(4)
        cols[0].metric("💰 Сумма", f"{amount:,.0f} ₽".replace(",", " "))
        cols[1].metric("⏱️ Срок", f"{zr.deadline_days} дн." if zr.deadline_days else "—")
        cols[2].metric("👤 Ответственный", zr.responsible[:12] if zr.responsible else "—")
        cols[3].metric("⚖️ ЮД", "Да" if zr.requires_legal else "Нет")
        
        if zr.approval_route:
            st.info("📍 " + " → ".join(zr.approval_route))
        for w in zr.warnings:
            st.warning(w)
    
    st.markdown("---")
    st.markdown("### 📄 Текст договора")
    
    # Загрузка файлов
    uploaded = st.file_uploader("Загрузить файл (TXT, DOCX, PDF)", type=["txt", "docx", "pdf"])
    if uploaded:
        success, result = DocumentLoader.load_file(uploaded)
        if success:
            st.session_state.contract_text = SecurityValidator.sanitize_text(result)
            st.success(f"✅ Загружено: {len(st.session_state.contract_text):,} символов")
        else:
            st.error(result)
    
    contract_text = st.text_area("Или вставьте текст:", value=st.session_state.contract_text, height=200)
    if contract_text != st.session_state.contract_text:
        st.session_state.contract_text = SecurityValidator.sanitize_text(contract_text)
    
    # Выбор типовой формы
    all_tf = {**BUILTIN_TYPICAL_FORMS, **st.session_state.custom_typical_forms}
    tf_options = [("Не сравнивать", None)] + [(f"{v['name']} ({v['code']})", k) for k, v in all_tf.items()]
    selected_tf_key = st.selectbox("Сравнить с типовой формой:", tf_options, format_func=lambda x: x[0])[1]
    
    col_an1, col_an2 = st.columns(2)
    with col_an1:
        if st.button("🔍 Сравнить с ТФ и найти риски", type="primary", use_container_width=True):
            if len(st.session_state.contract_text) < 100:
                st.error("Минимум 100 символов")
            elif not selected_tf_key:
                st.warning("Выберите типовую форму")
            else:
                with st.spinner("Анализируем..."):
                    tf = all_tf.get(selected_tf_key)
                    st.session_state.comparison_result = AdvancedComparator.full_comparison(st.session_state.contract_text, tf)
                    st.session_state.selected_tf = tf
                st.rerun()
    
    with col_an2:
        api_key = st.session_state.get("_openai_key") or st.session_state.get("_anthropic_key")
        if st.button("🤖 AI-анализ рисков", use_container_width=True, disabled=not api_key):
            if not api_key:
                st.warning("Введите API-ключ в Настройках")
            elif not st.session_state.comparison_result:
                st.warning("Сначала выполните сравнение с ТФ")
            else:
                with st.spinner("AI анализирует..."):
                    prompt = AIAnalyzer.generate_risk_analysis_prompt(st.session_state.contract_text, st.session_state.comparison_result)
                    if st.session_state.get("_openai_key"):
                        st.session_state.ai_analysis = AIAnalyzer.call_openai(prompt, st.session_state._openai_key)
                    else:
                        st.session_state.ai_analysis = AIAnalyzer.call_anthropic(prompt, st.session_state._anthropic_key)
                st.rerun()
    
    # Результаты сравнения
    if st.session_state.comparison_result:
        render_comparison_result(st.session_state.comparison_result)
    
    # AI анализ
    if st.session_state.ai_analysis:
        st.markdown("---")
        st.markdown("### 🤖 Экспертное заключение (AI)")
        st.markdown(st.session_state.ai_analysis)

def render_comparison_result(result: Dict):
    st.markdown("---")
    st.markdown(f"### 📋 Сравнение с {result.get('form_name', 'ТФ')}")
    
    compliance = result.get('compliance_score', 0)
    
    cols = st.columns(4)
    cols[0].metric("📊 Соответствие ТФ", f"{compliance:.0f}%")
    cols[1].metric("✅ Найдено разделов", len(result.get("found_sections", [])))
    cols[2].metric("❌ Отсутствует", len(result.get("missing_sections", [])))
    cols[3].metric("⚠️ Отклонений", len(result.get("deviation_sections", [])))
    
    st.markdown(f"**Резюме:** {result.get('summary', '')}")
    
    red_risks = [r for r in result.get('risks', []) if r.get('risk_level') == 'red']
    yellow_risks = [r for r in result.get('risks', []) if r.get('risk_level') == 'yellow']
    
    if red_risks:
        with st.expander(f"🔴 КРИТИЧЕСКИЕ РИСКИ ({len(red_risks)})", expanded=True):
            for r in red_risks:
                st.markdown(f"""<div class="risk-item red"><strong>⚠️ {r.get('issue', '')}</strong><br><small>{r.get('context', '')[:300]}...</small></div>""", unsafe_allow_html=True)
    
    if yellow_risks:
        with st.expander(f"🟡 ПРЕДУПРЕЖДЕНИЯ ({len(yellow_risks)})"):
            for r in yellow_risks:
                st.markdown(f"""<div class="risk-item yellow"><strong>⚡ {r.get('issue', '')}</strong><br><small>{r.get('context', '')[:300]}...</small></div>""", unsafe_allow_html=True)
    
    with st.expander("📑 Детальный анализ по разделам"):
        for s in result.get('sections_analysis', []):
            comp = s.get('comparison', {})
            score = comp.get('combined_score', 0) * 100
            status = comp.get('status', 'missing')
            status_icon = {"match": "✅", "partial": "⚠️", "deviation": "❌", "missing": "❌"}.get(status, "❓")
            score_class = "score-high" if score >= 70 else "score-medium" if score >= 40 else "score-low"
            
            st.markdown(f"""**{status_icon} {s.get('section_name', '')}** {'(обязательный)' if s.get('required') else ''} <span class="section-score {score_class}">{score:.0f}%</span>""", unsafe_allow_html=True)
            
            if comp.get('matched_section_name'):
                st.caption(f"Найден как: {comp['matched_section_name']}")
            
            for r in s.get('risks', []):
                lvl = "🔴" if r.get('risk_level') == 'red' else "🟡"
                st.markdown(f"  {lvl} {r.get('issue', '')}")
            
            st.markdown("---")

def render_reports_tab():
    st.markdown("### 📊 Генерация отчётов")
    
    if not st.session_state.comparison_result and not st.session_state.zone_result:
        st.info("Сначала выполните анализ договора на вкладке «Анализ»")
        return
    
    # Сборка данных
    inp = st.session_state.get('current_input') or AnalysisInput()
    
    data = {
        "analysis_id": hashlib.md5(datetime.now().isoformat().encode()).hexdigest()[:8],
        "counterparty": inp.counterparty or "Не указан",
        "contract_number": inp.contract_number or "",
        "contract_date": str(inp.contract_date) if inp.contract_date else "",
        "contract_type": inp.document_type or "",
        "document_form": inp.document_form.value if inp.document_form else "",
        "amount": inp.amount or 0,
        "legal_status": inp.legal_status.value if inp.legal_status else "",
        "legal_status_label": LEGAL_STATUS_LABELS.get(inp.legal_status, "") if inp.legal_status else "",
    }
    
    if st.session_state.zone_result:
        zr = st.session_state.zone_result
        data.update({
            "zone": zr.zone.value if zr.zone else "",
            "zone_reason": zr.zone_reason or "",
            "requires_legal": zr.requires_legal,
            "deadline_days": zr.deadline_days or 0,
            "approver": zr.responsible or "",
        })
    
    if st.session_state.comparison_result:
        cr = st.session_state.comparison_result
        data.update({
            "compliance_score": cr.get("compliance_score", 0),
            "tf_name": cr.get("form_name", ""),
            "tf_code": cr.get("form_code", ""),
            "risks": cr.get("risks", []),
            "sections_analysis": cr.get("sections_analysis", []),
            "missing_sections": cr.get("missing_sections", []),
            "deviation_sections": cr.get("deviation_sections", []),
            "recommendations": cr.get("recommendations", []),
            "summary": cr.get("summary", ""),
        })
    
    # Риск-скор
    red_count = sum(1 for r in data.get('risks', []) if r.get('risk_level') == 'red')
    yellow_count = sum(1 for r in data.get('risks', []) if r.get('risk_level') == 'yellow')
    data["risk_score"] = min(10, red_count * 2 + yellow_count * 0.7)
    
    # Заключение
    zone = data.get("zone", "")
    if zone == "red" or data["risk_score"] >= 6:
        data["conclusion"] = "Договор содержит КРИТИЧЕСКИЕ РИСКИ. Рекомендуется направить на полную юридическую экспертизу и устранить выявленные отклонения до подписания."
    elif zone == "yellow" or data["risk_score"] >= 3:
        data["conclusion"] = "Договор требует внимания. Рекомендуется согласовать с ЮД и рассмотреть возможность корректировки отдельных условий."
    else:
        data["conclusion"] = f"Существенных рисков не выявлено. Соответствие типовой форме: {data.get('compliance_score', 0):.0f}%. Допускается подписание в установленном порядке."
    
    st.markdown("#### 📋 Предпросмотр")
    zone_emoji = {"green": "🟢", "yellow": "🟡", "red": "🔴"}.get(zone, "⚪")
    st.markdown(f"""
    **Контрагент:** {data.get('counterparty')}  
    **Сумма:** {data.get('amount', 0):,.0f} ₽  
    **Зона:** {zone_emoji} {zone.upper() if zone else 'НЕ ОПРЕДЕЛЕНА'}  
    **Риск-скор:** {data['risk_score']:.1f}/10  
    **Соответствие ТФ:** {data.get('compliance_score', 0):.0f}%
    """)
    
    st.markdown("---")
    st.markdown("#### 📥 Скачать отчёты")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # PDF отчет
        try:
            pdf_data = ReportGenerator.generate_pdf_report(data, st.session_state.user, st.session_state.ai_analysis)
            st.download_button(
                "📥 PDF отчёт",
                pdf_data,
                f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "application/pdf",
                use_container_width=True
            )
        except:
            st.warning("PDF недоступен, используйте TXT")
    
    with col2:
        # DOCX отчет
        try:
            docx_data = ReportGenerator.generate_docx_report(data, st.session_state.user, st.session_state.ai_analysis)
            st.download_button(
                "📥 DOCX отчёт",
                docx_data,
                f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except:
            st.warning("DOCX недоступен")
    
    with col3:
        # TXT отчет
        txt_report = ReportGenerator.generate_text_report(data, st.session_state.user, st.session_state.ai_analysis)
        st.download_button(
            "📥 TXT отчёт",
            txt_report.encode('utf-8'),
            f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            use_container_width=True
        )
    
    st.markdown("---")
    st.markdown("#### 📥 JSON отчёты")
    
    col4, col5, col6 = st.columns(3)
    
    with col4:
        # JSON для 1С
        json_1c = ReportGenerator.generate_json_for_1c(data, st.session_state.user)
        st.download_button(
            "📥 JSON для 1С",
            json_1c,
            f"1c_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            use_container_width=True
        )
    
    with col5:
        # JSON для базы знаний
        json_kb = ReportGenerator.generate_json_knowledge_base(data, st.session_state.user, st.session_state.contract_text)
        st.download_button(
            "📥 JSON База знаний",
            json_kb,
            f"kb_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            use_container_width=True
        )
    
    with col6:
        if st.button("💾 В историю", use_container_width=True):
            add_to_history(data)
            st.success("✅ Сохранено!")
    
    # Просмотр
    with st.expander("👁️ Просмотр текстового отчёта"):
        st.text(txt_report)

def render_typical_forms_tab():
    st.markdown("### 📂 Типовые формы")
    
    st.markdown("#### ⬆️ Загрузить свою типовую форму")
    
    uploaded_tf = st.file_uploader("Загрузить ТФ (JSON)", type=["json"], key="tf_upload")
    if uploaded_tf:
        try:
            tf_data = json.loads(uploaded_tf.read().decode('utf-8'))
            if "name" in tf_data and "sections" in tf_data:
                code = tf_data.get("code", f"USER-{len(st.session_state.custom_typical_forms)+1}")
                st.session_state.custom_typical_forms[code] = tf_data
                st.success(f"✅ Загружена: {tf_data['name']}")
            else:
                st.error("Неверный формат")
        except Exception as e:
            st.error(f"Ошибка: {e}")
    
    st.markdown("##### Формат JSON:")
    st.code("""{
  "name": "Название",
  "code": "ТФ-001",
  "version": "1.0",
  "sections": {
    "1. РАЗДЕЛ": {
      "required": true,
      "template": "Текст...",
      "keywords": ["ключ1", "ключ2"],
      "risk_patterns": [{"pattern": "regex", "risk": "red", "issue": "Описание"}]
    }
  },
  "global_risk_patterns": []
}""", language="json")
    
    st.markdown("---")
    st.markdown("#### 📋 Доступные типовые формы")
    
    all_tf = {**BUILTIN_TYPICAL_FORMS, **st.session_state.custom_typical_forms}
    
    for key, tf in all_tf.items():
        is_custom = key in st.session_state.custom_typical_forms
        badge = "🔵 Пользовательская" if is_custom else "⚪ Встроенная"
        
        with st.expander(f"📄 {tf.get('name', key)} ({tf.get('code', key)}) {badge}"):
            st.markdown(f"**Версия:** {tf.get('version', '1.0')} | **Дата:** {tf.get('date', '')}")
            
            st.markdown("**Разделы:**")
            for sname, sdata in tf.get("sections", {}).items():
                req = "✅" if sdata.get("required") else "⚪"
                st.markdown(f"{req} **{sname}**")
            
            st.download_button(
                "📥 Скачать JSON",
                json.dumps(tf, ensure_ascii=False, indent=2),
                f"{key}.json",
                "application/json",
                key=f"dl_{key}"
            )
            
            if is_custom:
                if st.button(f"🗑️ Удалить", key=f"del_{key}"):
                    del st.session_state.custom_typical_forms[key]
                    st.rerun()

def render_regulation_tab():
    st.markdown("""
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 12px; color: white; margin-bottom: 1rem;">
        <h3 style="margin: 0;">📋 Регламент взаимодействия с ЮД</h3>
        <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">АО «НПК» | Версия 3.0</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="zone-card green"><div class="zone-title">🟢 ЗЕЛЕНАЯ — Самостоятельно</div><p>ТФ ≤100К | Нетиповые ≤50К</p></div>', unsafe_allow_html=True)
    st.markdown('<div class="zone-card yellow"><div class="zone-title">🟡 ЖЕЛТАЯ — ЮД 5 дней</div><p>ТФ 100К-5М | Нетип. >50К | Ед. поставщик >100К | Договор >3 лет</p></div>', unsafe_allow_html=True)
    st.markdown('<div class="zone-card red"><div class="zone-title">🔴 КРАСНАЯ — Полное сопровождение</div><p>>5М | Тендеры >3М | Вагоны/локомотивы | ВЭД | ПО | Недвижимость | Претензии | Госорганы</p></div>', unsafe_allow_html=True)
    st.warning("⚠️ ЗАПРЕЩЕНО подписание документов Желтой/Красной зоны без визы ЮД!")

def render_history_tab():
    st.markdown("### 📜 История анализов")
    if not st.session_state.history:
        st.info("История пуста")
        return
    
    for h in st.session_state.history:
        z = {"green": "🟢", "yellow": "🟡", "red": "🔴"}.get(h.get("zone"), "⚪")
        with st.expander(f"{z} {h.get('counterparty', 'N/A')} | {h.get('timestamp', '')[:10]}"):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**ID:** {h.get('id')}\n**Сумма:** {h.get('amount', 0):,.0f} ₽\n**Риск-скор:** {h.get('risk_score', 0):.1f}")
            with col2:
                st.markdown(f"**Соответствие ТФ:** {h.get('compliance_score', 0):.0f}%")
            st.download_button("📥 JSON", json.dumps(h, ensure_ascii=False, indent=2), f"analysis_{h.get('id')}.json", key=f"hist_{h.get('id')}")

def render_settings_tab():
    st.markdown("### ⚙️ Настройки")
    st.warning("API-ключи НЕ сохраняются между сессиями")
    
    with st.expander("🔑 API-ключи для AI-анализа"):
        openai_key = st.text_input("OpenAI API Key", type="password", placeholder="sk-...", value=st.session_state.get("_openai_key", ""))
        anthropic_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...", value=st.session_state.get("_anthropic_key", ""))
        if st.button("Применить"):
            if openai_key:
                st.session_state._openai_key = openai_key
            if anthropic_key:
                st.session_state._anthropic_key = anthropic_key
            st.success("✅ Ключи применены")
    
    st.markdown("---")
    st.markdown("""
    **Legal Traffic Light v5.1 Enterprise**  
    © АО «НПК» 2025
    
    **Возможности:**
    - Загрузка: TXT, DOCX, PDF
    - Отчёты: PDF, DOCX, TXT, JSON
    - Алгоритм: Jaccard + TF-IDF + N-grams + Levenshtein
    - AI-анализ: OpenAI / Anthropic
    - Загрузка пользовательских ТФ
    
    **Зависимости:**
    ```
    pip install streamlit python-docx fpdf2 PyPDF2 requests
    ```
    """)

def render_main():
    st.markdown("""
    <div class="main-header">
        <div class="traffic-light"><span class="tl-red"></span><span class="tl-yellow"></span><span class="tl-green"></span></div>
        <h2 style="margin: 0;">🚦 Legal Traffic Light <span class="version-badge">v5.1 Enterprise</span></h2>
        <p style="margin: 0.5rem 0 0 0; opacity: 0.8;">Система анализа договоров АО «НПК»</p>
    </div>
    """, unsafe_allow_html=True)
    
    render_sidebar()
    
    tabs = st.tabs(["📝 Анализ", "📊 Отчёты", "📂 Типовые формы", "📋 Регламент", "📜 История", "⚙️ Настройки"])
    with tabs[0]:
        render_analysis_tab()
    with tabs[1]:
        render_reports_tab()
    with tabs[2]:
        render_typical_forms_tab()
    with tabs[3]:
        render_regulation_tab()
    with tabs[4]:
        render_history_tab()
    with tabs[5]:
        render_settings_tab()

def main():
    apply_styles()
    init_session()
    if not st.session_state.authenticated:
        render_auth()
    else:
        render_main()

if __name__ == "__main__":
    main()
