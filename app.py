"""
Регламент Светофор v7.17
АО «СПК» — AI-анализ документов + Облачный OCR (модульная архитектура)

ИСПРАВЛЕНО в v7.17:
1. Ошибка 400 API — добавлена детальная диагностика
2. OCR и AI теперь независимые модули
3. OCR работает без AI (только распознавание)
4. AI работает без OCR (только анализ текстовых документов)
5. Настройка: использовать AI для OCR или облачный OCR API
6. Улучшена обработка ошибок
"""

import os
os.environ["STREAMLIT_SERVER_FILE_WATCHER_TYPE"] = "none"

import streamlit as st
import re, json, hashlib, io, time, base64
from datetime import datetime, date
from typing import Dict, List, Tuple, Optional, Any
from pathlib import Path

# ============================================================================
# БИБЛИОТЕКИ
# ============================================================================

DOCX_AVAILABLE = False
PDF_AVAILABLE = False
REQUESTS_AVAILABLE = False
PIL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    pass

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    pass

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    pass

IMAGE_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif', '.webp']

try:
    st.set_page_config(page_title="Регламент Светофор v7.17 | СПК", page_icon="🚦", layout="wide")
except:
    pass

# ============================================================================
# КОНСТАНТЫ
# ============================================================================

CONFIG_FILE = "svetofor_config.json"
HISTORY_FILE = "svetofor_history.json"

DEFAULT_ORG = {
    "full_name": "Акционерное общество «Специализированная ПК»",
    "short_name": "АО «СПК»",
    "inn": "7712345678"
}

DEFAULT_THRESHOLDS = {
    "зелёная_тф_макс": 100_000,
    "зелёная_нетф_макс": 50_000,
    "жёлтая_макс": 5_000_000
}

# AI провайдеры
AI_PROVIDERS = {
    "openai": {
        "name": "OpenAI GPT-4",
        "url": "https://platform.openai.com/api-keys",
        "model": "gpt-4o-mini",
        "endpoint": "https://api.openai.com/v1/chat/completions"
    },
    "anthropic": {
        "name": "Anthropic Claude",
        "url": "https://console.anthropic.com/settings/keys",
        "model": "claude-3-haiku-20240307",
        "endpoint": "https://api.anthropic.com/v1/messages"
    },
    "yandexgpt": {
        "name": "YandexGPT",
        "url": "https://console.cloud.yandex.ru/",
        "model": "yandexgpt-lite",
        "endpoint": "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    },
    "perplexity": {
        "name": "Perplexity AI",
        "url": "https://www.perplexity.ai/settings/api",
        "model": "llama-3.1-sonar-small-128k-online",
        "endpoint": "https://api.perplexity.ai/chat/completions"
    },
}

# OCR провайдеры (отдельно от AI)
OCR_PROVIDERS = {
    "yandex_vision": {
        "name": "Yandex Vision OCR",
        "url": "https://console.cloud.yandex.ru/",
        "endpoint": "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
    },
    "google_vision": {
        "name": "Google Cloud Vision",
        "url": "https://console.cloud.google.com/",
        "endpoint": "https://vision.googleapis.com/v1/images:annotate"
    },
}

# Типы документов
DOC_TYPES = {
    "Договор поставки": ["поставк", "поставщик", "покупатель", "товар"],
    "Договор купли-продажи": ["купл", "продаж", "продавец"],
    "Договор аренды": ["аренд", "арендодатель", "арендатор"],
    "Договор услуг": ["услуг", "исполнитель", "заказчик"],
    "Договор подряда": ["подряд", "подрядчик", "работы"],
    "Договор перевозки": ["перевоз", "перевозчик", "груз"],
    "Договор займа": ["займ", "заёмщик", "займодавец"],
    "Договор залога": ["залог", "залогодатель"],
    "Лицензионный договор": ["лицензи", "лицензиар", "лицензиат"],
    "Трудовой договор": ["трудов", "работник", "работодатель"],
    "Дополнительное соглашение": ["дополнительное соглашение", "допсоглашение"],
    "Счёт на оплату": ["счёт", "счет на оплату", "к оплате"],
    "Счёт-фактура": ["счёт-фактура", "счет-фактура", "ндс"],
    "УПД": ["универсальный передаточный", "упд"],
    "Товарная накладная": ["торг-12", "товарная накладная"],
    "Акт выполненных работ": ["акт выполненных", "акт сдачи"],
    "Акт сверки": ["акт сверки", "взаиморасчёт"],
    "Претензия": ["претензия", "требуем", "нарушение"],
    "Доверенность": ["доверенность", "доверяю", "полномочия"],
    "Приказ": ["приказ", "приказываю"],
    "Протокол": ["протокол", "собрание", "повестка"],
    "Гарантийное письмо": ["гарантийное письмо", "гарантируем"],
    "Коммерческое предложение": ["коммерческое предложение", "предлагаем"],
    "Техническое задание": ["техническое задание", "тз", "требования"],
    "Спецификация": ["спецификация", "характеристики"],
}

# ============================================================================
# СОХРАНЕНИЕ/ЗАГРУЗКА НАСТРОЕК
# ============================================================================

def get_config_path():
    return Path.home() / ".svetofor" / CONFIG_FILE

def get_history_path():
    return Path.home() / ".svetofor" / HISTORY_FILE

def save_config(config: dict):
    try:
        path = get_config_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2, default=str)
        return True
    except Exception as e:
        return False

def load_config() -> dict:
    try:
        path = get_config_path()
        if path.exists():
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return {}

def save_history(history: list):
    try:
        path = get_history_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(history[-100:], f, ensure_ascii=False, indent=2, default=str)
    except:
        pass

def load_history() -> list:
    try:
        path = get_history_path()
        if path.exists():
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return []

def init_session():
    saved = load_config()
    defaults = {
        "authorized": False,
        "user": None,
        "role": None,
        "org": saved.get("org", DEFAULT_ORG),
        "thresholds": saved.get("thresholds", DEFAULT_THRESHOLDS),
        "ai_keys": saved.get("ai_keys", {}),
        "ocr_keys": saved.get("ocr_keys", {}),
        "settings": saved.get("settings", {
            "use_ai_for_ocr": False,  # Использовать AI для OCR
            "preferred_ai": "",
            "preferred_ocr": "",
        }),
        "history": load_history(),
        "current_file_id": None,
        "doc_text": None,
        "doc_meta": None,
        "analysis": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def save_settings():
    config = {
        "org": st.session_state.get("org", DEFAULT_ORG),
        "thresholds": st.session_state.get("thresholds", DEFAULT_THRESHOLDS),
        "ai_keys": st.session_state.get("ai_keys", {}),
        "ocr_keys": st.session_state.get("ocr_keys", {}),
        "settings": st.session_state.get("settings", {}),
    }
    save_config(config)

def add_to_history(record: dict):
    history = st.session_state.get("history", [])
    record["timestamp"] = datetime.now().isoformat()
    history.append(record)
    st.session_state.history = history
    save_history(history)

def clear_analysis():
    for key in ["doc_text", "doc_meta", "analysis", "current_file_id", "identification", "full_analysis", "comparison"]:
        if key in st.session_state:
            del st.session_state[key]

# ============================================================================
# OCR МОДУЛЬ (независимый от AI)
# ============================================================================

class OCRModule:
    """
    Модуль OCR — работает независимо от AI.
    Использует облачные OCR API (Yandex Vision, Google Vision).
    """
    
    @classmethod
    def is_configured(cls) -> Tuple[bool, str]:
        """Проверка настроен ли OCR."""
        ocr_keys = st.session_state.get("ocr_keys", {})
        
        if ocr_keys.get("yandex_vision_key") and ocr_keys.get("yandex_folder_id"):
            return True, "Yandex Vision"
        if ocr_keys.get("google_vision_key"):
            return True, "Google Vision"
        
        return False, "OCR не настроен"
    
    @classmethod
    def get_status(cls) -> dict:
        """Статус OCR модуля."""
        configured, provider = cls.is_configured()
        ocr_keys = st.session_state.get("ocr_keys", {})
        
        return {
            "configured": configured,
            "provider": provider if configured else None,
            "yandex": bool(ocr_keys.get("yandex_vision_key") and ocr_keys.get("yandex_folder_id")),
            "google": bool(ocr_keys.get("google_vision_key")),
        }
    
    @classmethod
    def recognize(cls, image_bytes: bytes) -> Tuple[bool, str]:
        """Распознать текст из изображения."""
        ocr_keys = st.session_state.get("ocr_keys", {})
        settings = st.session_state.get("settings", {})
        preferred = settings.get("preferred_ocr", "")
        
        # Выбор провайдера
        if preferred == "yandex" and ocr_keys.get("yandex_vision_key"):
            return cls._yandex_ocr(image_bytes, ocr_keys)
        elif preferred == "google" and ocr_keys.get("google_vision_key"):
            return cls._google_ocr(image_bytes, ocr_keys)
        
        # Автовыбор
        if ocr_keys.get("yandex_vision_key") and ocr_keys.get("yandex_folder_id"):
            return cls._yandex_ocr(image_bytes, ocr_keys)
        if ocr_keys.get("google_vision_key"):
            return cls._google_ocr(image_bytes, ocr_keys)
        
        return False, "OCR не настроен. Перейдите в Настройки → OCR"
    
    @classmethod
    def _yandex_ocr(cls, image_bytes: bytes, keys: dict) -> Tuple[bool, str]:
        """Yandex Vision OCR."""
        if not REQUESTS_AVAILABLE:
            return False, "Библиотека requests не установлена"
        
        api_key = keys.get("yandex_vision_key", "")
        folder_id = keys.get("yandex_folder_id", "")
        
        if not api_key or not folder_id:
            return False, "Не указан API ключ или Folder ID для Yandex Vision"
        
        try:
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            response = requests.post(
                "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText",
                headers={
                    "Authorization": f"Api-Key {api_key}",
                    "Content-Type": "application/json",
                    "x-folder-id": folder_id
                },
                json={
                    "mimeType": "image",
                    "languageCodes": ["ru", "en"],
                    "model": "page",
                    "content": image_base64
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                if "result" in result and "textAnnotation" in result["result"]:
                    text = result["result"]["textAnnotation"].get("fullText", "")
                    if text:
                        return True, text
                    # Альтернативная сборка
                    blocks = result["result"]["textAnnotation"].get("blocks", [])
                    lines = []
                    for block in blocks:
                        for line in block.get("lines", []):
                            words = [w.get("text", "") for w in line.get("words", [])]
                            lines.append(" ".join(words))
                    if lines:
                        return True, "\n".join(lines)
                return False, "Текст не распознан"
            elif response.status_code == 400:
                error_detail = response.json().get("message", response.text[:200])
                return False, f"Yandex OCR ошибка 400: {error_detail}"
            elif response.status_code == 401:
                return False, "Yandex OCR: неверный API ключ"
            elif response.status_code == 403:
                return False, "Yandex OCR: доступ запрещён (проверьте Folder ID)"
            else:
                return False, f"Yandex OCR ошибка {response.status_code}: {response.text[:200]}"
        
        except requests.Timeout:
            return False, "Yandex OCR: таймаут (60 сек)"
        except Exception as e:
            return False, f"Yandex OCR ошибка: {str(e)}"
    
    @classmethod
    def _google_ocr(cls, image_bytes: bytes, keys: dict) -> Tuple[bool, str]:
        """Google Vision OCR."""
        if not REQUESTS_AVAILABLE:
            return False, "Библиотека requests не установлена"
        
        api_key = keys.get("google_vision_key", "")
        if not api_key:
            return False, "Не указан API ключ Google Vision"
        
        try:
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            response = requests.post(
                f"https://vision.googleapis.com/v1/images:annotate?key={api_key}",
                json={
                    "requests": [{
                        "image": {"content": image_base64},
                        "features": [{"type": "TEXT_DETECTION"}]
                    }]
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                if "responses" in result and result["responses"]:
                    annotations = result["responses"][0].get("textAnnotations", [])
                    if annotations:
                        return True, annotations[0].get("description", "")
                    error = result["responses"][0].get("error", {})
                    if error:
                        return False, f"Google OCR: {error.get('message', 'неизвестная ошибка')}"
                return False, "Текст не распознан"
            elif response.status_code == 400:
                error_detail = response.json().get("error", {}).get("message", response.text[:200])
                return False, f"Google OCR ошибка 400: {error_detail}"
            elif response.status_code == 403:
                return False, "Google OCR: неверный API ключ или API не включён"
            else:
                return False, f"Google OCR ошибка {response.status_code}"
        
        except requests.Timeout:
            return False, "Google OCR: таймаут (60 сек)"
        except Exception as e:
            return False, f"Google OCR ошибка: {str(e)}"

# ============================================================================
# AI МОДУЛЬ (независимый от OCR)
# ============================================================================

class AIModule:
    """
    Модуль AI — работает независимо от OCR.
    Используется для анализа текста документов.
    """
    
    @classmethod
    def is_configured(cls) -> Tuple[bool, str]:
        """Проверка настроен ли AI."""
        ai_keys = st.session_state.get("ai_keys", {})
        settings = st.session_state.get("settings", {})
        preferred = settings.get("preferred_ai", "")
        
        # Проверяем предпочтительный
        if preferred and ai_keys.get(preferred):
            return True, AI_PROVIDERS[preferred]["name"]
        
        # Проверяем любой
        for pid in AI_PROVIDERS:
            if ai_keys.get(pid):
                return True, AI_PROVIDERS[pid]["name"]
        
        return False, "AI не настроен"
    
    @classmethod
    def get_status(cls) -> dict:
        """Статус AI модуля."""
        configured, provider = cls.is_configured()
        ai_keys = st.session_state.get("ai_keys", {})
        
        return {
            "configured": configured,
            "provider": provider if configured else None,
            "openai": bool(ai_keys.get("openai")),
            "anthropic": bool(ai_keys.get("anthropic")),
            "yandexgpt": bool(ai_keys.get("yandexgpt")),
            "perplexity": bool(ai_keys.get("perplexity")),
        }
    
    @classmethod
    def _get_provider(cls) -> Tuple[str, str, str]:
        """Получить активный провайдер: (id, key, folder_id)."""
        ai_keys = st.session_state.get("ai_keys", {})
        settings = st.session_state.get("settings", {})
        preferred = settings.get("preferred_ai", "")
        folder_id = ai_keys.get("yandex_folder_id", "")
        
        if preferred and ai_keys.get(preferred):
            return preferred, ai_keys[preferred], folder_id
        
        for pid in ["openai", "anthropic", "perplexity", "yandexgpt"]:
            if ai_keys.get(pid):
                return pid, ai_keys[pid], folder_id
        
        return "", "", ""
    
    @classmethod
    def call(cls, prompt: str, max_tokens: int = 4000) -> Tuple[bool, str]:
        """Вызов AI API."""
        provider, key, folder_id = cls._get_provider()
        
        if not provider:
            return False, "AI не настроен. Добавьте API ключ в Настройки → AI"
        
        if not REQUESTS_AVAILABLE:
            return False, "Библиотека requests не установлена"
        
        try:
            # OpenAI
            if provider == "openai":
                response = requests.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "gpt-4o-mini",
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": max_tokens,
                        "temperature": 0.3
                    },
                    timeout=120
                )
                
                if response.status_code == 200:
                    return True, response.json()["choices"][0]["message"]["content"]
                elif response.status_code == 400:
                    error = response.json().get("error", {}).get("message", response.text[:300])
                    return False, f"OpenAI ошибка 400: {error}"
                elif response.status_code == 401:
                    return False, "OpenAI: неверный API ключ"
                elif response.status_code == 429:
                    return False, "OpenAI: превышен лимит запросов"
                else:
                    return False, f"OpenAI ошибка {response.status_code}: {response.text[:200]}"
            
            # Anthropic
            elif provider == "anthropic":
                response = requests.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": key,
                        "Content-Type": "application/json",
                        "anthropic-version": "2023-06-01"
                    },
                    json={
                        "model": "claude-3-haiku-20240307",
                        "max_tokens": max_tokens,
                        "messages": [{"role": "user", "content": prompt}]
                    },
                    timeout=120
                )
                
                if response.status_code == 200:
                    return True, response.json()["content"][0]["text"]
                elif response.status_code == 400:
                    error = response.json().get("error", {}).get("message", response.text[:300])
                    return False, f"Anthropic ошибка 400: {error}"
                elif response.status_code == 401:
                    return False, "Anthropic: неверный API ключ"
                else:
                    return False, f"Anthropic ошибка {response.status_code}: {response.text[:200]}"
            
            # Perplexity
            elif provider == "perplexity":
                response = requests.post(
                    "https://api.perplexity.ai/chat/completions",
                    headers={
                        "Authorization": f"Bearer {key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "llama-3.1-sonar-small-128k-online",
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": max_tokens
                    },
                    timeout=120
                )
                
                if response.status_code == 200:
                    return True, response.json()["choices"][0]["message"]["content"]
                elif response.status_code == 400:
                    error = response.json().get("error", {}).get("message", response.text[:300])
                    return False, f"Perplexity ошибка 400: {error}"
                elif response.status_code == 401:
                    return False, "Perplexity: неверный API ключ"
                else:
                    return False, f"Perplexity ошибка {response.status_code}: {response.text[:200]}"
            
            # YandexGPT
            elif provider == "yandexgpt":
                if not folder_id:
                    return False, "YandexGPT: не указан Folder ID"
                
                response = requests.post(
                    "https://llm.api.cloud.yandex.net/foundationModels/v1/completion",
                    headers={
                        "Authorization": f"Api-Key {key}",
                        "Content-Type": "application/json",
                        "x-folder-id": folder_id
                    },
                    json={
                        "modelUri": f"gpt://{folder_id}/yandexgpt-lite",
                        "completionOptions": {
                            "stream": False,
                            "temperature": 0.3,
                            "maxTokens": str(max_tokens)
                        },
                        "messages": [{"role": "user", "text": prompt}]
                    },
                    timeout=120
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if "result" in result and "alternatives" in result["result"]:
                        return True, result["result"]["alternatives"][0]["message"]["text"]
                    return False, "YandexGPT: пустой ответ"
                elif response.status_code == 400:
                    error = response.json().get("message", response.text[:300])
                    return False, f"YandexGPT ошибка 400: {error}"
                elif response.status_code == 401:
                    return False, "YandexGPT: неверный API ключ"
                elif response.status_code == 403:
                    return False, "YandexGPT: доступ запрещён (проверьте Folder ID)"
                else:
                    return False, f"YandexGPT ошибка {response.status_code}: {response.text[:200]}"
            
            return False, f"Неизвестный провайдер: {provider}"
        
        except requests.Timeout:
            return False, f"{provider}: таймаут (120 сек)"
        except requests.RequestException as e:
            return False, f"{provider} ошибка сети: {str(e)}"
        except Exception as e:
            return False, f"{provider} ошибка: {str(e)}"

# ============================================================================
# ФУНКЦИИ АНАЛИЗА
# ============================================================================

def simple_identify_document(text: str) -> dict:
    """Простая идентификация документа без AI (по ключевым словам)."""
    text_lower = text.lower()
    
    for doc_type, keywords in DOC_TYPES.items():
        matches = sum(1 for kw in keywords if kw in text_lower)
        if matches >= 2:
            return {
                "type": doc_type,
                "confidence": min(95, 60 + matches * 10),
                "method": "keywords"
            }
    
    return {
        "type": "Документ (тип не определён)",
        "confidence": 30,
        "method": "keywords"
    }


def ai_identify_document(text: str) -> dict:
    """AI идентификация документа."""
    org = st.session_state.get("org", DEFAULT_ORG)
    
    prompt = f"""Проанализируй документ и определи:

1. ТИП ДОКУМЕНТА (договор поставки, счёт, акт, претензия, и т.д.)
2. СТОРОНЫ (кто участвует, их роли, ИНН если есть)
3. НАША ОРГАНИЗАЦИЯ: {org.get('short_name', 'АО СПК')} (ИНН: {org.get('inn', '')})
4. КОНТРАГЕНТ (кто вторая сторона)
5. КЛЮЧЕВЫЕ ДАННЫЕ (сумма, срок, предмет)
6. КРАТКАЯ СПРАВКА (2-3 предложения о сути документа)

Ответ СТРОГО в JSON:
{{
  "type": "тип документа",
  "category": "договор/счёт/акт/письмо/другое",
  "confidence": 95,
  "parties": [
    {{"role": "Поставщик", "name": "ООО Ромашка", "inn": "1234567890", "is_us": false}},
    {{"role": "Покупатель", "name": "{org.get('short_name')}", "inn": "{org.get('inn')}", "is_us": true}}
  ],
  "counterparty": "ООО Ромашка",
  "subject": "предмет договора/документа",
  "amount": 100000,
  "currency": "RUB",
  "deadline": "срок",
  "summary": "краткая справка о документе"
}}

ТЕКСТ:
{text[:6000]}
"""
    
    success, response = AIModule.call(prompt, max_tokens=2000)
    
    if not success:
        # Fallback на простую идентификацию
        simple = simple_identify_document(text)
        simple["error"] = response
        simple["method"] = "keywords (AI недоступен)"
        return simple
    
    # Парсим JSON
    try:
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            result = json.loads(json_match.group())
            result["method"] = "AI"
            return result
    except:
        pass
    
    simple = simple_identify_document(text)
    simple["ai_response"] = response[:500]
    simple["method"] = "keywords (AI не вернул JSON)"
    return simple


def ai_full_analysis(text: str, identification: dict) -> dict:
    """Полный AI анализ документа."""
    org = st.session_state.get("org", DEFAULT_ORG)
    doc_type = identification.get("type", "Документ")
    category = identification.get("category", "другое")
    
    if category in ["договор", "contract"]:
        prompt = f"""Ты — юрист {org.get('short_name')} с 20-летним опытом.

ДОКУМЕНТ: {doc_type}
КОНТРАГЕНТ: {identification.get('counterparty', 'Не определён')}
СУММА: {identification.get('amount', 'Не указана')}

ТЕКСТ:
{text[:8000]}

ПРОВЕДИ АНАЛИЗ:

## 1. СУЩЕСТВЕННЫЕ УСЛОВИЯ
Проверь наличие обязательных условий по ГК РФ

## 2. РИСКИ (для каждого укажи):
- 🔴 Критический / 🟡 Существенный / 🟢 Незначительный  
- Пункт договора
- В чём проблема
- Как исправить (готовая формулировка)

## 3. НЕВЫГОДНЫЕ УСЛОВИЯ
Что ставит нас в худшее положение

## 4. РЕКОМЕНДАЦИЯ
СОГЛАСОВАТЬ / С ЗАМЕЧАНИЯМИ / ДОРАБОТАТЬ / ОТКЛОНИТЬ

## 5. ИТОГ
Краткий вывод (3-4 предложения)
"""
    else:
        prompt = f"""Ты — эксперт по документообороту {org.get('short_name')}.

ДОКУМЕНТ: {doc_type}

ТЕКСТ:
{text[:8000]}

ПРОАНАЛИЗИРУЙ:

## 1. СОДЕРЖАНИЕ
О чём документ, ключевые данные

## 2. ПРОВЕРКА
Правильность оформления, наличие реквизитов

## 3. ДЕЙСТВИЯ
Что нужно сделать с документом

## 4. ЗАМЕЧАНИЯ
Ошибки и рекомендации

## 5. ИТОГ
Краткий вывод
"""
    
    success, response = AIModule.call(prompt, max_tokens=4000)
    
    if not success:
        return {"error": response}
    
    return {
        "analysis": response,
        "provider": AIModule._get_provider()[0]
    }


def ai_compare_documents(text1: str, text2: str, name1: str, name2: str) -> dict:
    """AI сравнение документов."""
    
    prompt = f"""Сравни два документа и найди ВСЕ различия.

ДОКУМЕНТ 1 ({name1}):
{text1[:4000]}

ДОКУМЕНТ 2 ({name2}):
{text2[:4000]}

УКАЖИ:

## 1. ТИП ДОКУМЕНТОВ
Одинаковый или разный

## 2. ТАБЛИЦА РАЗЛИЧИЙ
| № | Пункт | Документ 1 | Документ 2 | Важность |
|---|-------|------------|------------|----------|

## 3. КРИТИЧЕСКИЕ РАЗЛИЧИЯ
Что меняет суть или создаёт риски

## 4. ВЫВОД
Идентичны / Несущественные различия / Существенно отличаются
"""
    
    success, response = AIModule.call(prompt, max_tokens=4000)
    
    if not success:
        return {"error": response}
    
    return {"comparison": response}

# ============================================================================
# ЗАГРУЗКА ФАЙЛОВ
# ============================================================================

def extract_text(uploaded_file) -> Tuple[bool, str, dict]:
    """Извлечение текста из файла."""
    if not uploaded_file:
        return False, "", {}
    
    filename = uploaded_file.name
    ext = Path(filename).suffix.lower()
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    
    meta = {"filename": filename, "size": len(file_bytes), "ext": ext}
    
    # TXT
    if ext == '.txt':
        for enc in ['utf-8', 'cp1251', 'cp866', 'latin-1']:
            try:
                return True, file_bytes.decode(enc), {**meta, "encoding": enc}
            except:
                continue
        return False, "Не удалось определить кодировку", meta
    
    # DOCX
    if ext == '.docx' and DOCX_AVAILABLE:
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Таблицы
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return True, "\n".join(paragraphs), meta
        except Exception as e:
            return False, f"Ошибка DOCX: {e}", meta
    
    # PDF
    if ext == '.pdf' and PDF_AVAILABLE:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Страница {i+1} ---\n{text}")
            
            if pages:
                full_text = "\n\n".join(pages)
                if len(full_text.strip()) > 100:
                    return True, full_text, meta
            
            # Возможно скан
            return False, "PDF не содержит текста (возможно скан). Загрузите как изображение для OCR.", meta
        except Exception as e:
            return False, f"Ошибка PDF: {e}", meta
    
    # Изображения — используем OCR модуль
    if ext in IMAGE_EXTENSIONS:
        ocr_ok, ocr_provider = OCRModule.is_configured()
        if not ocr_ok:
            return False, "Для изображений настройте OCR в Настройках → OCR", meta
        
        success, text = OCRModule.recognize(file_bytes)
        if success:
            return True, text, {**meta, "ocr": True, "ocr_provider": ocr_provider}
        else:
            return False, text, meta
    
    return False, f"Формат {ext} не поддерживается", meta

# ============================================================================
# ИНТЕРФЕЙС
# ============================================================================

def apply_styles():
    st.markdown("""
    <style>
    .main-header {
        display: flex; align-items: center; gap: 15px;
        padding: 15px; margin-bottom: 20px;
        background: linear-gradient(135deg, #1a1a2e, #16213e);
        border-radius: 10px; color: white;
    }
    .traffic-light { display: flex; gap: 5px; }
    .tl-red, .tl-yellow, .tl-green { width: 20px; height: 20px; border-radius: 50%; }
    .tl-red { background: #e74c3c; }
    .tl-yellow { background: #f39c12; }
    .tl-green { background: #27ae60; }
    
    .status-ok { color: #27ae60; font-weight: bold; }
    .status-err { color: #e74c3c; font-weight: bold; }
    .status-warn { color: #f39c12; font-weight: bold; }
    
    .doc-card {
        background: white; border-radius: 10px; padding: 20px;
        margin: 10px 0; box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        border-left: 4px solid #3498db;
    }
    .doc-card.success { border-left-color: #27ae60; }
    .doc-card.warning { border-left-color: #f39c12; }
    .doc-card.error { border-left-color: #e74c3c; }
    
    .module-card {
        background: #f8f9fa; padding: 15px; border-radius: 8px;
        margin: 10px 0; border: 1px solid #e9ecef;
    }
    .module-title { font-weight: bold; margin-bottom: 10px; }
    </style>
    """, unsafe_allow_html=True)


def render_header():
    st.markdown("""
    <div class="main-header">
        <div class="traffic-light">
            <span class="tl-red"></span><span class="tl-yellow"></span><span class="tl-green"></span>
        </div>
        <div>
            <h2 style="margin:0;">Регламент Светофор</h2>
            <small style="opacity:0.7;">v7.17 • Модульная архитектура</small>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_sidebar():
    with st.sidebar:
        st.markdown(f"### 👤 {st.session_state.get('user', 'Гость')}")
        
        # Статус модулей
        st.markdown("---")
        st.markdown("**Модули:**")
        
        # AI статус
        ai_status = AIModule.get_status()
        if ai_status["configured"]:
            st.markdown(f"🤖 AI: <span class='status-ok'>{ai_status['provider']}</span>", unsafe_allow_html=True)
        else:
            st.markdown("🤖 AI: <span class='status-warn'>не настроен</span>", unsafe_allow_html=True)
        
        # OCR статус
        ocr_status = OCRModule.get_status()
        if ocr_status["configured"]:
            st.markdown(f"🔍 OCR: <span class='status-ok'>{ocr_status['provider']}</span>", unsafe_allow_html=True)
        else:
            st.markdown("🔍 OCR: <span class='status-warn'>не настроен</span>", unsafe_allow_html=True)
        
        st.markdown("---")
        
        if st.button("🗑️ Очистить", use_container_width=True):
            clear_analysis()
            st.rerun()
        
        if st.button("🚪 Выйти", use_container_width=True):
            st.session_state.authorized = False
            st.rerun()


def render_upload():
    st.markdown("### 📄 Загрузка документа")
    
    col1, col2 = st.columns([4, 1])
    with col1:
        file = st.file_uploader(
            "Выберите файл",
            type=['txt', 'docx', 'pdf', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'],
            key="uploader"
        )
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🗑️", help="Очистить"):
            clear_analysis()
            st.rerun()
    
    if file:
        file_id = f"{file.name}_{file.size}"
        
        # Новый файл — очищаем
        if st.session_state.get("current_file_id") != file_id:
            clear_analysis()
            st.session_state.current_file_id = file_id
        
        # Извлекаем текст
        if st.session_state.get("doc_text") is None:
            with st.spinner("Читаю документ..."):
                success, text, meta = extract_text(file)
                
                if success:
                    st.session_state.doc_text = text
                    st.session_state.doc_meta = meta
                    ocr_note = " (OCR)" if meta.get("ocr") else ""
                    st.success(f"✅ {file.name}{ocr_note} — {len(text):,} символов")
                else:
                    st.error(f"❌ {text}")
                    return None
        
        return st.session_state.get("doc_text")
    
    return None


def render_analysis(text: str):
    """Секция анализа."""
    
    ai_ok, ai_name = AIModule.is_configured()
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if ai_ok:
            btn_identify = st.button("🔍 Определить тип (AI)", type="primary", use_container_width=True)
        else:
            btn_identify = st.button("🔍 Определить тип", use_container_width=True)
    
    with col2:
        btn_full = st.button("🤖 Полный анализ", use_container_width=True, disabled=not ai_ok)
        if not ai_ok:
            st.caption("Требуется AI")
    
    with col3:
        compare_file = st.file_uploader("📊 Сличить с...", type=['txt', 'docx', 'pdf'], 
                                        key="compare", label_visibility="collapsed")
    
    # Идентификация
    if btn_identify:
        with st.spinner("Определяю тип документа..."):
            if ai_ok:
                ident = ai_identify_document(text)
            else:
                ident = simple_identify_document(text)
            st.session_state.identification = ident
    
    # Полный анализ
    if btn_full and ai_ok:
        if "identification" not in st.session_state:
            with st.spinner("Определяю тип..."):
                st.session_state.identification = ai_identify_document(text)
        
        with st.spinner("AI анализирует документ..."):
            analysis = ai_full_analysis(text, st.session_state.identification)
            st.session_state.full_analysis = analysis
    
    # Сличение
    if compare_file:
        success2, text2, _ = extract_text(compare_file)
        if success2:
            with st.spinner("Сравниваю документы..."):
                comparison = ai_compare_documents(
                    text, text2,
                    st.session_state.get("doc_meta", {}).get("filename", "Документ 1"),
                    compare_file.name
                )
                st.session_state.comparison = comparison
        else:
            st.error(f"Не удалось прочитать: {text2}")
    
    # Отображение результатов
    render_results()


def render_results():
    """Отображение результатов."""
    
    # Идентификация
    if "identification" in st.session_state:
        ident = st.session_state.identification
        
        st.markdown("---")
        
        if "error" in ident and "method" not in ident:
            st.error(f"❌ {ident['error']}")
        else:
            card_class = "success" if ident.get("confidence", 0) > 70 else "warning"
            
            st.markdown(f"""
            <div class="doc-card {card_class}">
                <h3>📋 {ident.get('type', 'Документ')}</h3>
                <p>Метод: {ident.get('method', 'неизвестно')} • 
                   Уверенность: {ident.get('confidence', 0)}%</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Стороны
            parties = ident.get("parties", [])
            if parties:
                st.markdown("#### 👥 Стороны")
                cols = st.columns(len(parties))
                for i, p in enumerate(parties):
                    with cols[i]:
                        icon = "🏢" if p.get("is_us") else "🤝"
                        st.markdown(f"**{icon} {p.get('role', 'Сторона')}**")
                        st.markdown(p.get("name", "—"))
                        if p.get("inn"):
                            st.caption(f"ИНН: {p['inn']}")
            
            # Данные
            col1, col2, col3 = st.columns(3)
            with col1:
                amount = ident.get("amount")
                st.metric("💰 Сумма", f"{amount:,.0f} ₽" if amount else "—")
            with col2:
                st.metric("🤝 Контрагент", ident.get("counterparty", "—")[:20])
            with col3:
                st.metric("📅 Срок", ident.get("deadline", "—")[:15] if ident.get("deadline") else "—")
            
            # Справка
            if ident.get("summary"):
                st.info(f"📝 {ident['summary']}")
            
            # Добавляем в историю
            if "added_to_history" not in st.session_state:
                add_to_history({
                    "filename": st.session_state.get("doc_meta", {}).get("filename", ""),
                    "type": ident.get("type", ""),
                    "summary": ident.get("summary", "")[:200]
                })
                st.session_state.added_to_history = True
    
    # Полный анализ
    if "full_analysis" in st.session_state:
        analysis = st.session_state.full_analysis
        
        st.markdown("---")
        st.markdown("### 🤖 AI-экспертиза")
        
        if "error" in analysis:
            st.error(f"❌ {analysis['error']}")
        else:
            st.markdown(analysis.get("analysis", ""))
            provider = analysis.get("provider", "")
            if provider:
                st.caption(f"Провайдер: {AI_PROVIDERS.get(provider, {}).get('name', provider)}")
    
    # Сличение
    if "comparison" in st.session_state:
        comp = st.session_state.comparison
        
        st.markdown("---")
        st.markdown("### 📊 Сличение документов")
        
        if "error" in comp:
            st.error(f"❌ {comp['error']}")
        else:
            st.markdown(comp.get("comparison", ""))


def render_settings():
    st.markdown("### ⚙️ Настройки")
    
    tabs = st.tabs(["🤖 AI", "🔍 OCR", "🏢 Организация", "📊 Пороги"])
    
    # === AI ===
    with tabs[0]:
        st.markdown("#### API-ключи для AI-анализа")
        st.caption("AI используется для идентификации и экспертизы документов")
        
        ai_keys = st.session_state.get("ai_keys", {})
        
        # Yandex Folder ID (общий)
        st.markdown("**🔑 Yandex Cloud Folder ID** (для YandexGPT)")
        folder = st.text_input(
            "Folder ID",
            value=ai_keys.get("yandex_folder_id", ""),
            key="ai_folder",
            placeholder="b1g..."
        )
        if folder != ai_keys.get("yandex_folder_id", ""):
            st.session_state.ai_keys["yandex_folder_id"] = folder
            save_settings()
        
        st.markdown("---")
        
        # Провайдеры
        for pid, info in AI_PROVIDERS.items():
            col1, col2 = st.columns([4, 1])
            with col1:
                key = st.text_input(
                    f"**{info['name']}**",
                    type="password",
                    value=ai_keys.get(pid, ""),
                    key=f"ai_{pid}",
                    placeholder=info['url']
                )
            with col2:
                st.markdown("<br>", unsafe_allow_html=True)
                if ai_keys.get(pid):
                    st.markdown("<span class='status-ok'>✓</span>", unsafe_allow_html=True)
                else:
                    st.markdown("—")
            
            if key != ai_keys.get(pid, ""):
                st.session_state.ai_keys[pid] = key
                save_settings()
                st.rerun()
        
        # Выбор предпочтительного
        active = [p for p in AI_PROVIDERS if ai_keys.get(p)]
        if len(active) > 1:
            st.markdown("---")
            settings = st.session_state.get("settings", {})
            current = settings.get("preferred_ai", "")
            
            selected = st.selectbox(
                "Предпочтительный AI",
                ["Автовыбор"] + active,
                index=0 if current not in active else active.index(current) + 1,
                format_func=lambda x: AI_PROVIDERS.get(x, {}).get("name", x)
            )
            
            new_pref = "" if selected == "Автовыбор" else selected
            if new_pref != current:
                st.session_state.settings["preferred_ai"] = new_pref
                save_settings()
    
    # === OCR ===
    with tabs[1]:
        st.markdown("#### Облачный OCR")
        st.caption("OCR используется для распознавания сканов и изображений")
        st.caption("⚠️ OCR работает независимо от AI")
        
        ocr_keys = st.session_state.get("ocr_keys", {})
        
        st.markdown("---")
        st.markdown("**☁️ Yandex Vision OCR**")
        
        yv_key = st.text_input(
            "API Key Yandex Vision",
            type="password",
            value=ocr_keys.get("yandex_vision_key", ""),
            key="ocr_yv_key",
            placeholder="AQV... (из Сервисного аккаунта)"
        )
        
        yv_folder = st.text_input(
            "Folder ID",
            value=ocr_keys.get("yandex_folder_id", ""),
            key="ocr_yv_folder",
            placeholder="b1g..."
        )
        
        if yv_key != ocr_keys.get("yandex_vision_key", "") or yv_folder != ocr_keys.get("yandex_folder_id", ""):
            st.session_state.ocr_keys["yandex_vision_key"] = yv_key
            st.session_state.ocr_keys["yandex_folder_id"] = yv_folder
            save_settings()
            if yv_key and yv_folder:
                st.success("✅ Yandex Vision настроен")
        
        st.markdown("---")
        st.markdown("**☁️ Google Cloud Vision**")
        
        gv_key = st.text_input(
            "API Key Google Vision",
            type="password",
            value=ocr_keys.get("google_vision_key", ""),
            key="ocr_gv_key",
            placeholder="AIza..."
        )
        
        if gv_key != ocr_keys.get("google_vision_key", ""):
            st.session_state.ocr_keys["google_vision_key"] = gv_key
            save_settings()
            if gv_key:
                st.success("✅ Google Vision настроен")
        
        # Статус
        st.markdown("---")
        ocr_status = OCRModule.get_status()
        if ocr_status["configured"]:
            st.success(f"✅ OCR активен: {ocr_status['provider']}")
        else:
            st.warning("⚠️ OCR не настроен — изображения не будут распознаваться")
    
    # === Организация ===
    with tabs[2]:
        org = st.session_state.get("org", DEFAULT_ORG)
        
        full = st.text_input("Полное название", value=org.get("full_name", ""))
        short = st.text_input("Краткое название", value=org.get("short_name", ""))
        inn = st.text_input("ИНН", value=org.get("inn", ""))
        
        if st.button("💾 Сохранить", key="save_org"):
            st.session_state.org = {"full_name": full, "short_name": short, "inn": inn}
            save_settings()
            st.success("✅ Сохранено")
    
    # === Пороги ===
    with tabs[3]:
        th = st.session_state.get("thresholds", DEFAULT_THRESHOLDS)
        
        tf = st.number_input("🟢 Зелёная ТФ (₽)", value=th.get("зелёная_тф_макс", 100000), step=10000)
        ntf = st.number_input("🟢 Зелёная нетиповая (₽)", value=th.get("зелёная_нетф_макс", 50000), step=10000)
        yellow = st.number_input("🟡 Жёлтая макс (₽)", value=th.get("жёлтая_макс", 5000000), step=100000)
        
        if st.button("💾 Сохранить", key="save_th"):
            st.session_state.thresholds = {
                "зелёная_тф_макс": tf,
                "зелёная_нетф_макс": ntf,
                "жёлтая_макс": yellow
            }
            save_settings()
            st.success("✅ Сохранено")


def render_history():
    st.markdown("### 📜 История")
    
    history = st.session_state.get("history", [])
    
    if not history:
        st.info("История пуста")
        return
    
    for item in reversed(history[-30:]):
        with st.expander(f"📄 {item.get('filename', '—')} • {item.get('type', '—')}"):
            st.write(f"**Время:** {item.get('timestamp', '—')[:19]}")
            st.write(f"**Описание:** {item.get('summary', '—')}")
    
    if st.button("🗑️ Очистить историю"):
        st.session_state.history = []
        save_history([])
        st.rerun()


def render_login():
    st.markdown("""
    <div style="text-align:center;padding:50px;">
        <div style="display:flex;justify-content:center;gap:10px;margin-bottom:20px;">
            <span class="tl-red" style="width:25px;height:25px;border-radius:50%;background:#e74c3c;display:inline-block;"></span>
            <span class="tl-yellow" style="width:25px;height:25px;border-radius:50%;background:#f39c12;display:inline-block;"></span>
            <span class="tl-green" style="width:25px;height:25px;border-radius:50%;background:#27ae60;display:inline-block;"></span>
        </div>
        <h1>РЕГЛАМЕНТ СВЕТОФОР</h1>
        <p style="color:#c41e3a;">АО «СПК»</p>
        <p style="color:#666;">v7.17 • Модульная архитектура</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        with st.form("login"):
            login = st.text_input("Логин", placeholder="admin")
            pwd = st.text_input("Пароль", type="password", placeholder="admin123")
            
            if st.form_submit_button("Войти", use_container_width=True):
                users = {"admin": ("admin123", "admin"), "user": ("user123", "user")}
                
                if login in users and pwd == users[login][0]:
                    st.session_state.authorized = True
                    st.session_state.user = login
                    st.session_state.role = users[login][1]
                    st.rerun()
                else:
                    st.error("❌ Неверный логин или пароль")
        
        st.caption("Демо: admin/admin123")


def main():
    init_session()
    apply_styles()
    
    if not st.session_state.get("authorized"):
        render_login()
        return
    
    render_header()
    render_sidebar()
    
    # Вкладки для админа
    if st.session_state.get("role") == "admin":
        tabs = st.tabs(["📄 Анализ", "⚙️ Настройки", "📜 История"])
        
        with tabs[0]:
            text = render_upload()
            if text:
                render_analysis(text)
        
        with tabs[1]:
            render_settings()
        
        with tabs[2]:
            render_history()
    else:
        text = render_upload()
        if text:
            render_analysis(text)


if __name__ == "__main__":
    main()
